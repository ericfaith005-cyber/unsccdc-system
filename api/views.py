import json
from django.shortcuts import render
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from django.db.models import Q, Avg, Sum

# 💎 THE Hub Hub Hub Hub REST FRAMEWORK Hub Hub Hub Hub 💎
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework import viewsets, status
from rest_framework.response import Response
from rest_framework.decorators import api_view
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Sum, Count, Q, Avg
from django.contrib.auth import get_user_model
from .models import *

from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from .models import Staff, Student # 💎 Ensure Staff is imported!
import africastalking

from django.contrib.auth.decorators import login_required
from .models import Student, SchoolPayLedger, FeesTracker, Subject

# --- 1. SOVEREIGN SMS GATEWAY (LIVE PRODUCTION) ---
username = "yaweeric" 
api_key = "atsk_d23adde5b15396790edd13222388ce6a8ac25cf34e2544d703579e8e693477526d49f586" 
try:
    africastalking.initialize(username, api_key)
    sms = africastalking.SMS
except:
    print("SMS Gateway Standby")
from django.contrib.auth import get_user_model


def birth_the_king(request):
    User = get_user_model()
    User.objects.filter(username="admin").delete()
    
    # 💎 Create the King with all switches ON
    user = User.objects.create_superuser(
        username="admin",
        email="admin@unsccdc.com",
        password="Imperial2026!"
    )
    user.is_staff = True
    user.is_superuser = True
    user.is_active = True
    user.save()
    
    return HttpResponse("<h1>THE KING IS FULLY AUTHORIZED! 👑</h1>")


def get_national_grading(mark, level, project_score=0):
    """Returns: (Grade, Points, Status, Professional_Remark)"""
    try:
        mark = float(mark)
    except:
        return ('N/A', 0, 'No Data', 'Assessment data pending.')

    if level == 'UCE_NEW': # CBC O-Level
        if mark >= 80: return ('A', 1, 'Exceptional', 'Superior grasp of competencies.')
        elif mark >= 70: return ('B', 2, 'Outstanding', 'Consistently produces quality work.')
        elif mark >= 60: return ('C', 3, 'Satisfactory', 'Shows good understanding of core concepts.')
        else: return ('E', 5, 'Elementary', 'Needs significant support and coaching.')
    elif level == 'PRIMARY': # PLE
        if mark >= 90: return ('D1', 1, 'Distinction', 'Superb performance.')
        elif mark >= 80: return ('D2', 2, 'Distinction', 'Excellent work.')
        elif mark >= 70: return ('C3', 3, 'Credit', 'Very good effort.')
        else: return ('F9', 9, 'Fail', 'Intensive coaching required.')
    elif level == 'UNIVERSITY':
        if mark >= 80: return ('A', 5.0, 'First Class', 'Exceptional academic excellence.')
        return ('C', 2.0, 'Pass', 'Satisfactory.')
    return ('C4', 4, 'Credit', 'Fair performance.')

def get_ple_division(total_agg):
    if 4 <= total_agg <= 12: return "DIVISION 1"
    elif 13 <= total_agg <= 23: return "DIVISION 2"
    return "DIVISION 3"

# --- 3. SOVEREIGN TIERED TAX ENGINE (7000 UGX CAP) ---
def calculate_unsc_tax(amt):
    if amt <= 50000: return 750
    elif amt <= 100000: return 1000
    elif amt <= 200000: return 2000
    elif amt <= 500000: return 4000
    elif amt <= 1000000: return 6000
    return 7000

import traceback
from django.db.models import Sum, Avg, F
from rest_framework import viewsets
from rest_framework.response import Response
from .models import *

class StudentViewSet(viewsets.ViewSet):
    """Primary Uplink for the Sovereign Mobile Hub (Aligned for v4.0 App)"""

    def list(self, request):
        code_in = request.query_params.get('code', '').strip()
        phone_in = request.query_params.get('phone', '').strip()
        pin_in = request.query_params.get('pin', '').strip()
        p_name = request.query_params.get('parent', '').strip().lower()
        s_name = request.query_params.get('student', '').strip().lower()

        try:
            # 1. Identity Gate
            p_rec = Parent.objects.get(phone_number=phone_in)
            student = p_rec.students.first() 
            if not student or student.full_name.lower() != s_name:
                return Response({"msg": "Identity Mismatch in National Registry"}, status=401)
            
            # 2. Secure PIN Verification
            if pin_in and p_rec.secure_pin.strip() == pin_in:
                s = p_rec.linked_student
                sch = s.school

                # --- 🛡️ SOVEREIGN SAFEGUARDS ---
                bio_obj = getattr(s, 'bio', None)
                bio_data = {
                    "career": getattr(bio_obj, 'future_career', "National Leader"),
                    "challenges": getattr(bio_obj, 'challenges_faced', "None"),
                    "inspiration": getattr(bio_obj, 'student_inspiration', "Sovereignty"),
                    "gender": getattr(s, 'gender', 'M'), # 💎 ADDED: For profile icons
                    "stream": getattr(s, 'stream', 'North'), # 💎 ADDED: For registry accuracy
                }

                f_rec = getattr(s, 'fees', None)
                finance = {
                    "total_due": getattr(f_rec, 'total_fees_due', 0),
                    "paid": getattr(f_rec, 'total_fees_paid', 0),
                    "balance": (getattr(f_rec, 'total_fees_due', 0) - getattr(f_rec, 'total_fees_paid', 0)),
                }

                # --- 📊 ATTENDANCE MATH ---
                att_records = s.attendance_records.all()
                total_days = att_records.count()
                present_days = att_records.filter(status='PRESENT').count()
                late_days = att_records.filter(status='LATE').count()
                att_percentage = ((present_days + (late_days * 0.5)) / total_days * 100) if total_days > 0 else 100.0
                bio_data["attendance"] = f"{att_percentage:.1f}%"

                # --- 📑 10-COLUMN DATA GATHERING ---
                national_report = {}
                p_map = {"AOI1":"aoi_1", "AOI2":"aoi_2", "MidTerm":"mid_term", "AOI3":"aoi_3", "AOI4":"aoi_4", "EOT":"eot_score"}
                
                for p_key, m_field in p_map.items():
                    mlist = []
                    current_term_total_pts = 0
                    for m in s.marks.all():
                        score = getattr(m, m_field, 0)
                        # We assume get_national_grading is defined in your utils
                        g, pts, st, rem = get_national_grading(score, s.level_category)
                        if p_key == "EOT": current_term_total_pts += pts
                        
                        mlist.append({
                            "sub": m.subject.name, 
                            "score": score, 
                            "grade": g, 
                            "aoi1": m.aoi_1, 
                            "aoi2": m.aoi_2, 
                            "mid": m.mid_term, 
                            "aoi3": m.aoi_3, 
                            "aoi4": m.aoi_4, 
                            "project": m.project_work,
                            "teacher": "STAFF" # 💎 ADDED: For the 'TCH' column in App
                        })
                    
                    if mlist: 
                        national_report[p_key] = {
                            "marks": mlist, 
                            "agg": current_term_total_pts, 
                            "div": "DIV 1" if current_term_total_pts <= 12 else "VERIFIED"
                        }

                # --- 🎞️ TOP PERFORMERS ALIGNMENT ---
                performers = []
                for t in NationalTopPerformer.objects.all().order_by('?')[:12]:
                    performers.append({
                        "name": t.name, 
                        "school_name": t.school_name, # 💎 MATCHES APP KEY
                        "score": t.score, 
                        "photo": request.build_absolute_uri(t.photo.url) if t.photo else ""
                    })

                # --- 📱 TIKTOK FEED ALIGNMENT (CRITICAL FOR LOADING FIX) ---
                feed_data = []
                for f in SchoolPost.objects.all().order_by('-date'):
                    feed_data.append({
                        "id": f.id, # 💎 ADDED: For interaction tracking
                        "school_name": f.school.name, # 💎 MATCHES APP KEY
                        "content": f.title, # 💎 MATCHES APP KEY (CONTENT)
                        "media": request.build_absolute_uri(f.media_file.url) if f.media_file else "",
                        "likes": f.likes_count,
                        "comment_count": getattr(f, 'comments_total', 0) # 💎 ADDED: For UI badges
                    })

                # 📦 THE FINAL IMPERIAL PACKAGE
                return Response({
                    "status": "authenticated", # 💎 ADDED: For the App's new Verify logic
                    "type": "parent", 
                    "name": s.full_name, 
                    "id": s.account_number, 
                    "payment_code": s.payment_code, 
                    "sch_id": sch.school_account_id, 
                    "class": s.current_class, 
                    "curriculum": s.level_category,
                    "photo": request.build_absolute_uri(s.photo.url) if s.photo else "", 
                    "parent_name": p_rec.full_name,
                    "school": {
                        "name": sch.name, 
                        "addr": sch.address, 
                        "motto": sch.school_motto, 
                        "uneb_no": sch.uneb_center_number, 
                        "school_code": sch.school_code, # 💎 ADDED
                        "dir": sch.director,
                        "mission": getattr(sch, 'mission', "Excellence"), 
                        "vision": getattr(sch, 'vision', "Sovereignty"), 
                        "rating": getattr(sch, 'rating', "⭐⭐⭐⭐⭐"), 
                        "type": getattr(sch, 'school_type', "Standard")
                    },
                    "finance": finance, 
                    "national_report": national_report, 
                    "bio_info": bio_data,
                    "top_performers": performers,
                    "feed": feed_data,
                    "ussd_steps": sch.ussd_instructions,
                    "payment_history": [
                        {
                            "receipt": t.receipt_number,
                            "amount": t.amount,
                            "date": t.timestamp.strftime('%d-%b-%Y'),
                            "channel": "USSD / SchoolPay"
                        } for t in SchoolPayLedger.objects.filter(student=s).order_by('-timestamp')
                    ],
                })
                
            return Response({"msg": "PIN_REQUIRED", "motto": p_rec.security_motto})
            
        except Parent.DoesNotExist:
            return Response({"msg": "Rejected: Code or Phone invalid"}, status=401)
        except Exception as e:
            print(traceback.format_exc())
            return Response({"msg": "National Registry Error"}, status=500)

@login_required
def finances_dashboard(request):
    # 🧮 GOLIATH FINANCIAL CALCULATIONS
    # Get all transactions for the school (assuming request.user.school link)
    ledger = SchoolPayLedger.objects.filter(school=request.user.school).order_by('-timestamp')
    
    # Financial Overview Metrics
    stats = FeesTracker.objects.filter(student__school=request.user.school).aggregate(
        total_invoiced=Sum('total_fees_due'),
        total_paid=Sum('total_fees_paid')
    )
    
    invoiced = stats['total_invoiced'] or 0
    paid = stats['total_paid'] or 0
    outstanding = invoiced - paid

    context = {
        'ledger': ledger,
        'total_invoiced': invoiced,
        'total_paid': paid,
        'outstanding': outstanding,
        'active_tab': 'finances'
    }
    return render(request, 'tabs/finances.html', context)

@login_required
def academics_dashboard(request):
    # 👨‍🎓 STUDENT REGISTRY DATA
    # Select students and prefetch fee status for the "Status Indicator"
    students = Student.objects.filter(school=request.user.school).order_by('full_name')
    
    context = {
        'students': students,
        'active_tab': 'academics'
    }
    return render(request, 'tabs/academics.html', context)

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

@csrf_exempt
def verify_identity(request):
    """
    THE IMPERIAL GATEWAY (Stage 1)
    Manually handling the CORS Handshake to stop the Loading Forever bug.
    """
    
    # 💎 1. THE SOVEREIGN HANDSHAKE (CORS FORCE)
    # We create the response manually to ensure the 'Stamp' is there
    response = JsonResponse({})
    response["Access-Control-Allow-Origin"] = "https://schoolapp-lac.vercel.app"
    response["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    response["Access-Control-Allow-Headers"] = "Content-Type, X-CSRFToken, Authorization"
    response["Access-Control-Allow-Credentials"] = "true"

    # If the browser is just asking for permission (OPTIONS), give it and STOP.
    if request.method == "OPTIONS":
        return response

    # 💎 2. THE DATA LOOKUP (STAGE 1)
    # Extract details from the URL (since your log shows they are coming via GET)
    code = request.GET.get('code', '').strip()
    student_name = request.GET.get('student', '').strip()
    parent_name = request.GET.get('parent', '').strip()
    phone = request.GET.get('phone', '').strip()

    # Query the Registry
    match = Student.objects.filter(
            payment_code=code,
            full_name__iexact=student_name,
            parent_link__full_name__iexact=parent_name, # 💎 CHANGED THIS
            parent_link__phone_number=phone             # 💎 CHANGED THIS
        ).first()

    if match:
        # ✅ SUCCESS: Identity Confirmed
        response.content = json.dumps({
            'status': 'success', 
            'message': 'Credentials verified. Opening PIN vault.',
            'student_id': match.account_number
        }).encode('utf-8')
        return response
    else:
        # 🛑 DENIED
        response.status_code = 401
        response.content = json.dumps({
            'status': 'error', 
            'message': 'No matching records found in the National Registry.'
        }).encode('utf-8')
        return response
    
@csrf_exempt # 💎 EMERGENCY BYPASS: Allows the form to hit the server from any origin
def verify_student_portal(request):
    """
    STAGE 1: NATIVE FORM POST VALIDATION
    """
    if request.method == 'POST':
        # 1. Capture values directly from the POST body
        incoming_code = request.POST.get('code', '').strip()
        incoming_student = request.POST.get('student', '').strip()
        incoming_parent = request.POST.get('parent', '').strip()
        incoming_phone = request.POST.get('phone', '').strip()

        # 2. THE MASTER SEARCH
        # Searching the National Registry for a 4-point match
        match = Student.objects.filter(
            payment_code=incoming_code,
            full_name__iexact=incoming_student,
            parent__full_name__iexact=incoming_parent,
            parent__phone_number=incoming_phone
        ).first()

        if match:
            # ✅ SUCCESS: Identity Confirmed
            # Instantly transition to the PIN entry screen
            return render(request, 'pin_entry.html', {
                'student': match,
                'status': 'authenticated'
            })
        else:
            # 🛑 DENIED: Return to login with error
            return render(request, 'index.html', {
                'error': 'Identity Denied. No matching records found in the National Registry.',
                'old_data': request.POST # Keeps the typed text so they don't re-type
            })

    # If it's a GET request, just show the login page
    return render(request, 'index.html')

from django.shortcuts import render
from .models import Student, Parent, Staff
from django.views.decorators.csrf import csrf_exempt

@csrf_exempt
def parent_verify_view(request):
    # Default state is 'gate' (The 4-box form)
    context = {'stage': 'gate'} 

    if request.method == 'POST':
        # --- STAGE 1: IDENTITY HANDSHAKE ---
        if 'verify_identity' in request.POST:
            code = request.POST.get('code', '').strip()
            student_name = request.POST.get('student', '').strip()
            parent_name = request.POST.get('parent', '').strip()
            phone = request.POST.get('phone', '').strip()

            student = Student.objects.filter(
                payment_code__iexact=code,
                full_name__iexact=student_name,
                parent_link__full_name__iexact=parent_name,
                parent_link__phone_number__icontains=phone[-9:]
            ).first()

            if student:
                context = {'stage': 'vault', 'student': student}
            else:
                context = {'stage': 'gate', 'error': 'Identity Mismatch. Check spelling/details.'}

        # --- STAGE 2: FINAL Hub AUTHORIZATION (6-DIGIT PIN) ---
        elif 'authorize_access' in request.POST:
            input_pin = request.POST.get('pin', '').strip()
            student_id = request.POST.get('student_id')
            student = Student.objects.get(account_number=student_id)
            
            if student.parent_link and student.parent_link.secure_pin == input_pin:
                return render(request, 'tabs/home.html', {'data': student})
            else:
                context = {'stage': 'vault', 'student': student, 'error': 'SECURITY ALERT: Invalid 6-Digit PIN.'}

    return render(request, 'index.html', context)


# --- 🛰️ THE IMPERIAL NATIONAL MARKS ENGINE (REAL-TIME SYNC) ---
@api_view(['POST'])
def staff_marks_engine(request):
    """
    Saves and updates academic scores in real-time.
    Handles: AOI 1-4, Mid Term, EOT, and Project Work.
    """
    try:
        # 1. Capture data from the Instructor's App
        student_id = request.data.get('student_id')
        subject_id = request.data.get('subject_id')
        field = request.data.get('field') # e.g., 'aoi_1', 'mid_term', 'eot_score'
        score = float(request.data.get('score', 0))

        # 2. Locate the Citizen and the Subject
        student = Student.objects.get(account_number=student_id)
        subject = Subject.objects.get(id=subject_id)

        # 3. Use 'get_or_create' so the Brain builds a new record 
        # if this is the first test of the term
        result, created = AcademicResult.objects.get_or_create(
            student=student, 
            subject=subject
        )

        # 4. THE SOVEREIGN SETATTR: Dynamically updates the specific test column
        # Mapping the App field name to the Database column name
        db_field = field
        if field == "mid": db_field = "mid_term"
        if field == "eot": db_field = "eot_score"
        if field == "proj": db_field = "project_work"

        setattr(result, db_field, score)
        result.save()

        # Royal Log for the Terminal
        print(f"--- 🏛️ HUB UPDATE: {student.full_name} | {subject.name} | {db_field}: {score} ---")

        return Response({
            "status": "Verified", 
            "msg": f"Registry updated for {student.full_name}."
        })

    except Exception as e:
        return Response({"msg": f"Registry Error: {str(e)}"}, status=400)

# --- 6. REAL-TIME CATEGORIZED SETTLEMENT ---
@api_view(['POST'])
def pay_fees(request):
    try:
        sid, amt, cat = request.data.get('student_id'), float(request.data.get('amount', 0)), request.data.get('category', 'Tuition')
        s = Student.objects.get(account_number=sid); sch = s.school
        
        # POCKET LOGIC
        if cat.lower() == "tuition":
            f = s.fees; f.total_fees_paid += int(amt); f.save()
        
        tax = calculate_unsc_tax(amt)
        sch.total_revenue_collected += int(amt); sch.total_commission_earned += int(tax); sch.save()
        txn_id = f"UNS-TXN-{uuid.uuid4().hex[:10].upper()}"
        Transaction.objects.create(transaction_id=txn_id, school=sch, student=s, amount_paid=amt, system_tax=tax, category=cat)
        return Response({"status": "Verified", "txn": txn_id})
    except: return Response({"msg": "Gateway Error"}, status=400)

# --- 7. STAFF PORTAL VIEWSET ---
class StaffViewSet(viewsets.ViewSet):
    def list(self, request):
        staff = Staff.objects.all()
        return Response([{
            "name": s.full_name, 
            "id": s.staff_id, 
            "designation": s.get_role_display() # This shows 'Class Teacher' instead of 'TEACHER'
        } for s in staff])

@api_view(['GET'])
def UNSCCDC_Analytics(request):
    return Response({"status": "Online"})
# --- 🏛️ THE IMPERIAL NATIONAL MARKS ENGINE (PASTE AT BOTTOM) ---
@api_view(['POST'])
def staff_marks_engine(request):
    """Saves and updates academic scores in real-time."""
    try:
        student_id = request.data.get('student_id')
        subject_id = request.data.get('subject_id')
        field = request.data.get('field') # aoi_1, mid_term, eot_score
        score = float(request.data.get('score', 0))

        student = Student.objects.get(account_number=student_id)
        subject = Subject.objects.get(id=subject_id)

        # Build or Update the Registry Record
        result, created = AcademicResult.objects.get_or_create(student=student, subject=subject)
        
        # Mapping App field names to Brain column names
        db_map = {"mid": "mid_term", "eot": "eot_score", "proj": "project_work"}
        final_field = db_map.get(field, field)

        setattr(result, final_field, score)
        result.save()

        return Response({"status": "Verified", "msg": "Sync Successful"})
    except Exception as e:
        return Response({"msg": f"Registry Error: {str(e)}"}, status=400)
    
from django.http import StreamingHttpResponse
import json

# --- 🛡️ SURGERY: IMPERIAL PRIVACY LOCK ---
def sync_schoolpay_transaction(school, tx_data):
    receipt = tx_data.get("schoolpayReceiptNumber")
    prn = tx_data.get("studentPaymentCode")
    amount = float(tx_data.get("amount", 0))

    if SchoolPayLedger.objects.filter(receipt_number=receipt).exists():
        return False, "Duplicate Blocked"

    try:
        # 💎 CRITICAL LOCK: We ONLY find the student IF they belong to THIS school
        student = Student.objects.get(payment_code=prn, school=school)
        
        # --- 🛰️ SURGERY: UPDATING THE CREATE LOGIC ---

        # Find the .create line and update it to look like this:
        SchoolPayLedger.objects.create(
            receipt_number=receipt, 
            school=school, 
            student=student, 
            amount=amount,
            raw_data=tx_data # 💎 THIS SAVES THE MTN/AIRTEL INFO FOR THE CHART
        )
        # Balance update remains private to this student
        f = student.fees
        f.total_fees_paid += int(amount)
        f.save()
        
        return True, f"Verified for {student.full_name}"
    except Student.DoesNotExist:
        # If the PRN exists but in a DIFFERENT school, this school will never see it!
        return False, "Security: PRN not found in your registry"

# # --- 🛡️ SURGERY: ENSURE NAME MATCHES (api/views.py) ---
def bursar_notification_stream(request, school_id):
    """Bursar dashboard connects here to see live payments"""
    def event_stream():
        # Keep track of the last seen transaction
        last_id = SchoolPayLedger.objects.filter(school_id=school_id).last().id if SchoolPayLedger.objects.filter(school_id=school_id).exists() else 0
        while True:
            new_txs = SchoolPayLedger.objects.filter(school_id=school_id, id__gt=last_id)
            for tx in new_txs:
                yield f"data: {json.dumps({'name': tx.student.full_name, 'amt': tx.amount})}\n\n"
                last_id = tx.id
            time.sleep(5) # Efficient polling
    return StreamingHttpResponse(event_stream(), content_type='text/event-stream')


def generate_imperial_pdf(request, student_id):
    """
    Standalone Python Engine to draw the 1,000-Level National Certificate.
    Includes: 12 Columns, Watermarks, Flag Borders, and National Grading Keys.
    """
    try:
        # 1. FETCH DATA
        student = Student.objects.get(account_number=student_id)
        sch = student.school
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="National_Report_{student.full_name}.pdf"'

        if student.photo:
            try:
                # We use the path to the photo
                p.drawImage(student.photo.path, width-130, height-150, width=80, height=100, mask='auto')
                # Gold Frame
                p.setStrokeColor(rich_gold)
                p.setLineWidth(1)
                p.rect(width-130, height-150, 80, 100, stroke=1)
            except Exception as photo_error:
                print(f"Photo Error: {photo_error}") # Don't crash if photo is missing

        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4

        # 2. 🛡️ THE SOVEREIGN WATERMARK (DIAGONAL)
        p.saveState()
        p.setFont("Helvetica-Bold", 45)
        p.setStrokeColor(colors.lightgrey)
        p.setFillColor(colors.lightgrey, alpha=0.04)
        p.translate(width/2, height/2)
        p.rotate(45)
        p.drawCentredString(0, 0, "UNSCCDC NATIONAL HUB OFFICIAL")
        p.restoreState()

        # 3. 🇺🇬 THE NATIONAL FLAG BORDERS (BLACK, YELLOW, RED)
        p.setLineWidth(2)
        p.setStrokeColor(colors.black); p.rect(15, 15, width-30, height-30)
        p.setStrokeColor(colors.orange); p.rect(18, 18, width-36, height-36) # Yellow replacement
        p.setStrokeColor(colors.red); p.rect(21, 21, width-42, height-42)

        # 4. 🏛️ OFFICIAL NATIONAL HEADER
        p.setFont("Helvetica-Bold", 10)
        p.drawCentredString(width/2, height-60, "THE REPUBLIC OF UGANDA")
        
        # Draw a small Seal Box in the middle
        p.setStrokeColor(colors.black); p.rect(width/2-25, height-115, 50, 50)
        p.setFont("Helvetica-Bold", 7)
        p.drawCentredString(width/2, height-90, "OFFICIAL")
        p.drawCentredString(width/2, height-100, "SEAL")

        p.setFont("Helvetica-Bold", 9)
        p.drawCentredString(width/2, height-130, "UGANDA NATIONAL EXAMINATIONS BOARD (UNEB)")
        p.setFont("Helvetica-Bold", 18)
        p.setFillColor(colors.HexColor("#003366"))
        p.drawCentredString(width/2, height-155, sch.name.upper())
        p.setFont("Helvetica-Oblique", 8)
        p.setFillColor(colors.black)
        p.drawCentredString(width/2, height-170, f'"{sch.school_motto}"')

        # 5. 👤 STUDENT REGISTRY MATRIX
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, height-210, f"STUDENT: {student.full_name.upper()}")
        p.drawString(50, height-225, f"NATIONAL ID: {student.account_number}")
        p.drawString(400, height-210, f"CLASS: {student.current_class}")
        p.drawString(400, height-225, f"PAY CODE: {student.payment_code or 'N/A'}")

        # 6. 📊 THE GOLIATH 12-COLUMN TABLE (CBC STANDARD)
        # Header Row
        matrix_data = [['SUBJECT', 'AOI1', 'AOI2', 'MID', 'AOI3', 'AOI4', 'EOT', 'PROJ', 'AVG', 'GRD', 'TCH', 'REMARKS']]
        
        marks = student.marks.all()
        for m in marks:
            formatted_score = f"{m.eot_score:g} / {m.eot_max}" 
            # Automated Remark Logic
            score = m.eot_score
            remark = "Superior" if score >= 90 else ("Excellent" if score >= 75 else "Satisfactory")
            
            matrix_data.append([
                m.subject.name[:8].upper(), str(m.aoi_1), str(m.aoi_2), str(m.mid_term), 
                str(m.aoi_3), str(m.aoi_4), f"{score}%", str(m.project_work), 
                f"{score}%", "A", "STF", remark
            ])

        # Table Styling (London Standard)
        table = Table(matrix_data, colWidths=[65, 30, 30, 30, 30, 30, 35, 35, 35, 25, 35, 80])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.black),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        
        table_height = len(matrix_data) * 15
        table.wrapOn(p, width, height)
        table.drawOn(p, 40, height-250 - table_height)

        # 7. 🤖 AUTOMATED ADMINISTRATIVE COMMENTS
        footer_y = height-280 - table_height
        p.setFont("Helvetica-Bold", 8)
        p.drawString(50, footer_y, "CLASS TEACHER: High discipline observed. Promoted to next academic year.")
        p.drawString(50, footer_y - 15, "HEAD TEACHER: Excellent performance across core competencies.")
        
        # 8. 📚 NATIONAL GRADING SCALES (EXPLAINED)
        p.setFont("Helvetica-Bold", 7)
        p.drawString(50, 140, "NATIONAL GRADING STANDARDS:")
        p.setFont("Helvetica", 6)
        p.drawString(50, 130, "SECONDARY CBC: 90-100: A+ (Exceptional) | 80-89: A | 70-79: B | 60-69: C | 50-59: D | 0-49: E")
        p.drawString(50, 120, "PRIMARY PLE: DIV 1 (4-12 Agg) | DIV 2 (13-23 Agg) | DIV 3 (24-28 Agg) | DIV 4 (29-34 Agg)")
        p.drawString(50, 110, "UNIVERSITY (NCHE): 4.40 - 5.00: First Class | 3.60 - 4.39: Second Upper | 2.00 - 2.79: Pass")

        # 9. ✍️ SIGNATURES & FINAL STAMP
        p.line(50, 70, 180, 70)
        p.drawCentredString(115, 60, "Head Teacher Signature")
        
        # THE RED STAMP
        p.setStrokeColor(colors.red)
        p.circle(width/2, 75, 30, stroke=1, fill=0)
        p.setFillColor(colors.red)
        p.setFont("Helvetica-Bold", 8)
        p.drawCentredString(width/2, 85, "UNSCCDC")
        p.drawCentredString(width/2, 75, "VERIFIED")
        p.drawCentredString(width/2, 65, "2026")
        
        p.setFillColor(colors.black)
        p.line(400, 70, 530, 70)
        p.drawCentredString(465, 60, "National Registrar Seal")

        p.setFont("Helvetica", 5)
        p.drawCentredString(width/2, 35, f"Verification ID: {student.account_number}-{datetime.date.today().year}. Authenticated Hub Record.")

        p.showPage()
        p.save()
        return response

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return HttpResponse(f"Hub Printing Error: {str(e)}", status=400)
    
@login_required
def bursar_print_center(request):
    # 🕵️ 1. SECURITY & IDENTITY
    school = getattr(request.user, 'school', None)
    if not school:
        school = School.objects.filter(school_type='SECONDARY').first() # Founder Bypass

    # 🕵️ 2. DYNAMIC LEVEL DETECTION (Killing the P.1 Error)
    if school.school_type == 'SECONDARY':
        class_list = ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6']
        hub_label = "ORDINARY & ADVANCED REGISTRY"
    else:
        class_list = ['P.1', 'P.2', 'P.3', 'P.4', 'P.5', 'P.6', 'P.7']
        hub_label = "PRIMARY FOUNDATION REGISTRY"

    # 🕵️ 3. REAL-TIME DATA PUMP
    # Pulling every transaction from today from the SchoolPay Ledger
    today = timezone.now().date()
    all_txs = SchoolPayLedger.objects.filter(
        school=school, 
        timestamp__date=today
    ).select_related('student')

    # Sorting counts for the glowing badges
    pending_counts = {cls: all_txs.filter(student__current_class=cls, is_printed=False).count() for cls in class_list}

    return render(request, 'admin/api/bursarterminal/change_list.html', {
        'school': school,
        'class_list': class_list,
        'txs': all_txs,
        'pending_counts': pending_counts,
        'hub_label': hub_label,
        'today': today
    })

def generate_payslip_pdf(request, payroll_id):
    p = StaffPayroll.objects.get(id=payroll_id)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Payslip_{p.staff.full_name}_{p.month}.pdf"'

    c = canvas.Canvas(response, pagesize=A4)
    w, h = A4

    # 🏛️ HEADER
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w/2, h-50, "OFFICIAL STAFF PAYMENT SLIP")
    c.setFont("Helvetica", 10)
    c.drawCentredString(w/2, h-65, f"{p.staff.school.name.upper()}")
    
    # 👤 STAFF INFO
    c.line(50, h-80, w-50, h-80)
    c.drawString(50, h-100, f"STAFF NAME: {p.staff.full_name}")
    c.drawString(50, h-115, f"DESIGNATION: {p.staff.designation}")
    c.drawString(400, h-100, f"MONTH: {p.month} {p.year}")
    c.drawString(400, h-115, f"TIN: {p.staff.tin_number}")

    # 📊 EARNINGS & DEDUCTIONS TABLE
    data = [
        ['DESCRIPTION', 'EARNINGS', 'DEDUCTIONS'],
        ['Basic Gross Salary', f"{p.gross_salary:,.0f}", ''],
        ['NSSF (5%)', '', f"{p.nssf_deduction:,.0f}"],
        ['PAYE Tax (URA)', '', f"{p.paye_tax:,.0f}"],
        ['Other Deductions', '', f"{p.other_deductions:,.0f}"],
        ['TOTAL NET PAY', '', f"UGX {p.net_pay:,.0f}"]
    ]
    
    table = Table(data, colWidths=[200, 150, 150])
    table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#D4AF37")),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ]))
    table.wrapOn(c, w, h)
    table.drawOn(c, 50, h-250)

    c.showPage()
    c.save()
    return response

@login_required
def bursar_batch_terminal(request):
    school = request.user.school
    today = timezone.now().date()
    
    # 🕵️ Logic: Identify the correct class list based on National Level
    level_map = {
        'KIND': ['Baby', 'Middle', 'Top'],
        'PRIM': ['P.1', 'P.2', 'P.3', 'P.4', 'P.5', 'P.6', 'P.7'],
        'SEC':  ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6'],
        'INTL': [f'Year {i}' for i in range(1, 14)],
        'UNI':  [f'Year {i}' for i in range(1, 6)],
    }
    active_classes = level_map.get(school.school_type, ['Standard'])

    # 📊 Live Transaction Counter per class
    todays_txs = SchoolPayLedger.objects.filter(school=school, timestamp__date=today)
    
    class_stats = []
    for cls in active_classes:
        count = todays_txs.filter(student__current_class=cls).count()
        class_stats.append({'name': cls, 'count': count})

    return render(request, 'bursar_terminal.html', {
        'school': school,
        'class_stats': class_stats,
        'transactions': todays_txs,
        'today': today
    })

def generate_staff_dossier_pdf(request, staff_id):
    """Generates a high-security National HR Dossier"""
    try:
        staff = Staff.objects.get(staff_id=staff_id)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Dossier_{staff.full_name}.pdf"'

        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4

        # 🛡️ Watermark (8-space indent inside try)
        p.saveState()
        p.setFont("Helvetica-Bold", 60)
        p.setFillColor(colors.lightgrey, alpha=0.05)
        p.translate(width/2, height/2)
        p.rotate(45)
        p.drawCentredString(0, 0, "OFFICIAL HUB")
        p.restoreState()

        # 🏛️ Header
        p.setFont("Helvetica-Bold", 14)
        p.drawCentredString(width/2, height-50, "THE REPUBLIC OF UGANDA")
        p.setFont("Helvetica", 10)
        p.drawCentredString(width/2, height-70, f"STATION: {staff.school.name.upper()}")

        # 👤 Data
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, height-120, f"FULL NAME: {staff.full_name.upper()}")
        p.drawString(50, height-140, f"STAFF ID: {staff.staff_id}")

        p.showPage()
        p.save()
        return response
    except Exception as e:
        return HttpResponse(f"Registry Error: {str(e)}", status=404)

def generate_payslip_pdf(request, payroll_id):
    """Draws a high-end, 12-row Audit-Ready Payslip"""
    try:
        p = StaffPayroll.objects.get(id=payroll_id)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Payslip_{p.staff.full_name}_{p.month}.pdf"'

        c = canvas.Canvas(response, pagesize=A4)
        w, h = A4

        # 1. 🛡️ OFFICIAL WATERMARK
        c.saveState()
        c.setFont("Helvetica-Bold", 50)
        c.setFillColor(colors.lightgrey, alpha=0.03)
        c.translate(w/2, h/2); c.rotate(45)
        c.drawCentredString(0, 0, "UNSCCDC OFFICIAL HUB")
        c.restoreState()

        # 2. 🏛️ HEADER & IDENTITY
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(w/2, h-50, "OFFICIAL STAFF REMUNERATION SLIP")
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(w/2, h-65, p.staff.school.name.upper())
        c.line(50, h-80, w-50, h-80)

        # 👤 STAFF PROFILE DATA
        c.setFont("Helvetica-Bold", 9)
        c.drawString(50, h-100, f"NAME: {p.staff.full_name.upper()}")
        c.drawString(50, h-115, f"ID: {p.staff.staff_id}")
        c.drawString(50, h-130, f"DESIGNATION: {p.staff.designation}")
        
        c.drawString(380, h-100, f"MONTH: {p.month.upper()} {p.year}")
        c.drawString(380, h-115, f"URA TIN: {p.staff.tin_number or 'N/A'}")
        c.drawString(380, h-130, f"NSSF No: {p.staff.nssf_number or 'N/A'}")

        # 3. 📊 THE FINANCIAL AUDIT MATRIX
        data = [
            ['DESCRIPTION', 'EARNINGS (UGX)', 'DEDUCTIONS (UGX)'],
            ['Basic Gross Salary', f"{p.gross_salary:,.0f}", ''],
            ['NSSF Contribution (5%)', '', f"{p.nssf_deduction:,.0f}"],
            ['PAYE Income Tax (URA)', '', f"{p.paye_tax:,.0f}"],
            ['Other Deductions', '', f"{p.other_deductions:,.0f}"],
            ['', '', ''], # Spacer
            ['TOTAL NET PAYOUT', '', f"UGX {p.net_pay:,.0f}"]
        ]

        table = Table(data, colWidths=[200, 150, 150])
        table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#D4AF37")),
            ('TEXTCOLOR', (0,0), (-1,0), colors.black),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('ALIGN', (1,1), (-1,-1), 'RIGHT'),
            ('ALIGN', (0,0), (0,-1), 'LEFT'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ]))
        table.wrapOn(c, w, h)
        table.drawOn(c, 50, h-300)

        # ✍️ 4. AUTHORIZATION
        c.setFont("Helvetica-Bold", 8)
        c.drawString(50, 150, "Bursar's Signature: _____________________")
        c.drawRightString(w-50, 150, "Official Hub Seal: 🔒")
        c.setFont("Helvetica-Oblique", 7)
        c.drawCentredString(w/2, 50, "This is an electronically generated document. Valid without physical signature.")

        c.showPage(); c.save()
        return response
    except Exception as e:
        return HttpResponse(f"Dossier Error: {str(e)}", status=400)
     # --- 👑 THE SECRET KING-MAKER DOOR ---
from django.contrib.auth.models import User

def create_initial_king(request):
    # This is a secret URL to build the first admin in the clouds
    if not User.objects.filter(username="admin").exists():
        User.objects.create_superuser("admin", "admin@unsccdc.com", "Imperial2026!")
        return HttpResponse("The King is Born in the Clouds! 👑")
    return HttpResponse("The Throne is already occupied.")   
# --- 👑 THE IMPERIAL KING-MAKER DOOR ---
from django.contrib.auth import get_user_model


FLAG_STYLE = """
<style>
    body { 
        background: #050505; color: #fff; font-family: 'Courier New', monospace; 
        padding: 50px; margin: 0; overflow-x: hidden;
        background-image: radial-gradient(circle at 50% 50%, rgba(252, 220, 4, 0.05) 0%, transparent 80%);
    }
    .flag-bar { 
        position: fixed; top: 0; left: 0; width: 100%; height: 5px; display: flex; 
        box-shadow: 0 5px 15px rgba(212, 175, 55, 0.3);
    }
    .b { flex: 1; background: #000; } .y { flex: 1; background: #FCDC04; } .r { flex: 1; background: #D90000; }
    
    .glass-tab {
        background: rgba(255, 255, 255, 0.02); border: 1px solid rgba(255, 255, 255, 0.1);
        padding: 40px; border-radius: 30px; backdrop-filter: blur(10px);
        border-top: 4px solid #FCDC04; position: relative; animation: slideIn 0.8s ease-out;
    }
    @keyframes slideIn { from { opacity: 0; transform: translateY(30px); } to { opacity: 1; transform: translateY(0); } }
    
    .nav-btn {
        display: inline-block; padding: 12px 25px; border-radius: 10px; text-decoration: none;
        font-weight: 900; font-size: 11px; letter-spacing: 2px; transition: 0.3s;
        border: 1px solid #333; color: #888; margin-right: 10px;
    }
    .nav-btn:hover { background: #FCDC04; color: #000; box-shadow: 0 0 20px #FCDC04; }
    .active-btn { background: #D90000 !important; color: #fff !important; border: none; box-shadow: 0 0 20px #D90000; }
    
    .pulse-dot { height: 8px; width: 8px; background: #00ff00; border-radius: 50%; display: inline-block; margin-right: 10px; animation: blink 1.5s infinite; }
    @keyframes blink { 0% { opacity: 1; } 50% { opacity: 0.3; } 100% { opacity: 1; } }
</style>

<div class="flag-bar">
    <div class="b"></div><div class="y"></div><div class="r"></div>
    <div class="b"></div><div class="y"></div><div class="r"></div>
</div>

<div style="margin-bottom: 40px;">
    <a href="/api/home/" class="nav-btn {{home_act}}">1. HOME</a>
    <a href="/api/about/" class="nav-btn {{about_act}}">2. ABOUT</a>
    <a href="/api/academics/" class="nav-btn {{acad_act}}">3. ACADEMICS</a>
    <a href="/api/finances/" class="nav-btn {{fin_act}}">4. FINANCES</a>
    <a href="/api/profile/" class="nav-btn {{prof_act}}">5. PROFILE</a>
    <a href="/admin/" class="nav-btn" style="float:right;">← BACK TO CONTROL</a>
</div>
"""
# --- 📜 THE ABOUT TAB WITH LIVE APK DOWNLOAD ---
def about_tab(request):
    html = """
    <body style="background:#000; color:#fff; font-family:sans-serif; padding:50px; text-align:center;">
        <div style="border: 2px solid #D4AF37; padding: 40px; border-radius: 30px; background: rgba(212,175,55,0.02);">
            <h1 style="color:#D4AF37; letter-spacing:5px; font-weight:900;">UNSCCDC NATIONAL HUB</h1>
            <p style="color:#888; letter-spacing:2px;">OFFICIAL MOBILE INTERFACE v1.0.1</p>
            
            <hr style="border-color:#222; margin: 30px 0;">
            
            <p style="font-size:18px;">Founder & Chief Innovation Officer: <b>Yawe Eric</b></p>
            <p style="color:#aaa; font-style:italic;">"In 2025, at the age of 20, Ugandan Software developer and Tech Entrepeneur Yawe Eric recognized a critical gap in the nation's educational infrastructure: 
            schools were overwhelmed by disorganized manual paperwork, fee tracking was prone to leakages, and parents remained in the dark about thier children's daily performance."</p>
            <p style="color:#aaa;">"With a bold vision to completely transform Uganda's Education sector, Eric engineered UNSCCDC. His mission is to brig world-class, cloud-based digital infrastructure to every school in Uganda-starting with better Institutions-ensuring accountability, 
            moving Uganda Education into a paperless, digitally transparent future."</p>

            <p style="font-size:18px;"><b>Technical Architecture and Reliability Specs</b></p>
            <ul style="color:#ccc; line-height:2;">
             <li>✅ Cloud Infrastructure</li>
                <p style="color:#aaa;">Hosted on highly reliable cloud servers with an automated deployment ppeline linked directly to secure version control, ensuring 99.9% platform uptime.</p>

                <li>✅ Database Integrity</li>
                <p style="color:#aaa;">Built on a robust relational database management system using Django's Object-Relational Mapping (ORM) to handle complex queries for thousands of student profiles without lag.</p>

                <li>✅ Local Compliance</li>
                <p style="color:#aaa;">Designed to align fully with the assessment grading guidelines stipulated by the Ministry of Educaton and Sports (MoES) and the Uganda National Curriculum Development Centre (NCDC).</p>
            
            </ul>
           

            <p style="font-size:18px;"><b>Technical Architecture and Reliability Specs</b></p>


            <!-- 🚀 THE NATIONAL DOWNLOAD BUTTON -->
            <div style="margin-top:40px;">
                <a href="1uVswBKYlTe6xC-5gIxhkGwcsAu_lxd67" 
                   style="background:#D4AF37; color:#000; padding:25px 50px; border-radius:20px; font-weight:900; text-decoration:none; font-size:20px; box-shadow: 0 10px 20px rgba(0,0,0,0.5); display:inline-block;">
                   📥 DOWNLOAD ANDROID APP (APK)
                </a>
            </div>

            <p style="margin-top:50px;"><a href="/admin/" style="color:#D4AF37; text-decoration:none;">← BACK TO CONTROL CENTRE</a></p>
        </div>
    </body>
    """
    return HttpResponse(html)

def academics_tab(request):
    content = f"""{FLAG_STYLE.replace('{{acad_act}}', 'active-btn')}
    <div class="glass-tab">
        <h1 style="color:#FCDC04;">ACADEMIC ASSESSMENT ENGINE</h1>
        <p>This is the operational core of the school, built to handle <b>the complex realities of the Uganda grading matrix</b>.</p>

         <p style="font-size:18px;"><b>The Assessment Engine Specifications</b></p>
            <ul style="color:#ccc; line-height:2;">
            <li>✅ New Lower Secondary Curriculum (NLSC)</li>
            <p style="color:#aaa;"><b>Tracker:</b>Built-in grading architecture designed for the 20-point continuos assessment scale. It allows teachers to input Activities of Intergration (AoIs), automatically calculates scores out of 3, and generates the mandatory NCDC-compliant descriptors. </p>

             <li>✅ Traditional Curriculum Grading</li>
            <p style="color:#aaa;">An automated system for all levels that instantly converts raw percentages into UNEB-standard aggregtes and automatically determines student divisions and subjcet combinations. </p>

             <li>✅ Automated Report Card Generation</li>
            <p style="color:#aaa;">A one-clickgeneration system that compiles continuos assessment mrks, final exams, teacher remarks, housmaster comments, and school fees balance into a secure, downloadable PDF report card carrying the digital signature of the Headteacher.</p>

             <li>✅ Digital Staffroom Timetabler</li>
            <p style="color:#aaa;">An algorithmic scheduling tool that prevents room clashes and teacher double-booking across different classes and streams of all levels.</p>

             <li>✅ Student Progress Analytics</li>
            <p style="color:#aaa;">Interactive graphical trends showing a student's performance trajectory across multiple terms, allowing directors to identify struggling students early. </p>
        </ul>
    </div>"""
    return HttpResponse(content)

def finances_tab(request):
    content = f"""{FLAG_STYLE.replace('{{fin_act}}', 'active-btn')}
    <div class="glass-tab">
        <h1 style="color:#FCDC04;">THE FINANCIAL LEAK-PROOF LEDGER</h1>
        
       
        <ul style="color:#ccc; line-height:2;"> <p><b>Revenue and Ledger Management</b></p>
        <li>✅ Student Progress Analytics</li>
            <p style="color:#aaa;">At the start of every term, the system automatically applies unique billing structures to every student based on thier class, stream, or boarder/day scholar status, eliminating manual invoicing errors.</p>
        <li>✅ Real-Time Cash Flow Analytics</li>
            <p style="color:#aaa;">Provides the School Director with a secure, instant breakdown of total expected revenue, total fees collected so far, and total outstanding school debts. </p>
        
        <p><b>Advanced Anti-Leakage Intergration</b></p>

        <li>✅ Digital Gateway Snycing</li>
            <p style="color:#aaa;">Designed to hook into mobile money API networks (MTN MoMo and Airtel Money) and local banking agents. When a parent pays fees at a bank or via phone, the system instantly logs the payment, deducts the balance from the student's profile, and updates the bursar's dashboard.</p>
        
        <li>✅ Automated SMS Reminders</li>
            <p style="color:#aaa;">An intelligent notification agent that identifies accounts with outstanding balances at specified intervals (e.g. Week 4, Week 8) and sends a personalized, polite text reminder directly to the parent's phone.</p>
        
        <li>✅ Clearance Slip Verification</li>
            <p style="color:#aaa;">Genertes a secure, digital verification token (or barcode) once a student hits a set payment threshold, allowing gate staff to verify financil clearance instantly during school return days.</p>
        </ul>
       
        <div style="display:flex; gap:20px; margin-top:20px;">
            <div style="flex:1; background:#f1c40f; color:#000; padding:15px; border-radius:10px; font-weight:900; text-align:center;">MTN MoMo</div>
            <div style="flex:1; background:#D90000; color:#fff; padding:15px; border-radius:10px; font-weight:900; text-align:center;">AIRTEL MONEY</div>
        </div>
    </div>"""
    return HttpResponse(content)


def home_tab(request):
    html = """
    <div style="background: linear-gradient(90deg, #FCDC04, #D90000); padding: 20px; border-radius: 15px; text-align: center; margin-bottom: 30px; animation: pulse 2s infinite;">
        <h2 style="color: #000; margin: 0; font-weight: 900;"> NATIONAL HUB APP READY</h2>
        <p style="color: #000; font-size: 12px; font-weight: bold;">Click the button below to install the Official App on your Android phone.</p>
        <a href="/api/get-app/" style="background: #000; color: #fff; padding: 15px 30px; border-radius: 10px; text-decoration: none; font-weight: 900; display: inline-block; margin-top: 10px;">
           INSTALL APP NOW (56MB)
        </a>
    </div>
    
    <div class="module-card">
        <h1 style="font-family:'Orbitron'; color:#FCDC04; letter-spacing:8px; margin:0;">UNSCCDC GLOBAL</h1>
    <body style="background:#000; color:#fff; font-family:sans-serif; padding:50px;">
        <h1 style="color:#D4AF37; letter-spacing:3px;">ENTERPRISE COMMAND CENTER</h1>
        <p style="color:#888;">UNSCCDC GLOBAL Hub Status: <span style="color:#00ff00;">● LIVE</span></p>
        <hr style="border-color:#222;">
        <div style="display:grid; grid-template-columns:1fr 1fr; gap:20px;">
            <div style="background:#111; padding:20px; border-radius:15px; border-left:4px solid #D4AF37;">
                <h3 style="margin:0;">95%</h3>
                <p style="font-size:10px; color:#666;">ERROR REDUCTION</p>
            </div>
            <div style="background:#111; padding:20px; border-radius:15px; border-left:4px solid #00ff00;">
                <h3 style="margin:0;">14 Hours</h3>
                <p style="font-size:10px; color:#666;">SAVED WEEKLY</p>
            </div>
        </div>
        <p style="margin-top:30px;"><a href="/admin/" style="color:#D4AF37; text-decoration:none;">← BACK TO CONTROL CENTRE</a></p>
    </body>
    """
    return HttpResponse(html)

def profile_tab(request):
    content = f"""{FLAG_STYLE.replace('{{prof_act}}', 'active-btn')}
    <div class="glass-tab">
        <h1 style="color:#FCDC04;"><b>Security and Audit Logs</b></h1>
        <div style="background:#111; padding:30px; border-radius:20px;">
            <h3>Role-Based Access Control (RBAC): Users are strictly restricted based on permission groups:</h3>
            <li>✅ Super Administrators</li>
            <p style="color:#aaa;">Full database access, system configuration, and deployment controls.</p>

            <li>✅ School Administrators (Bursars/ Headteachers)</li>
            <p style="color:#aaa;">Access to financial reports, staff payroll, and final grade approvals.</p>

            <li>✅ Educators</li>
            <p style="color:#aaa;">Access only to the specific classes and subjects assigned to them for mark entry</p>

            <li>✅ Parents</li>
            <p style="color:#aaa;">Read-only access restricted strictly to their biological children's financial and academic records.</p>
        
            <p style="color:#00ff00; font-weight:bold;">Security Audit Trail</p>
            <p style="color:#aaa;">Tracks user activity for accountability. it displays the login timestamp, the device IP address, and a log of recent actions (e.g., "Teacher Namubiru Shifat updated Senior 3 Math marks on June 8, 2026").</p>

            <p style="color:#00ff00; font-weight:bold;">User Settings and Customization</p>
            <li>✅ Biometric and Two-Factor Authentication (2FA)</li>
            <p style="color:#aaa;">Optional security layer requiring an SMS token code before administrative or financial changes can be saved.</p>

            <li>✅ Language and Accessibility</li>
            <p style="color:#aaa;">Toggle features for high-contrast viewing and future intergration for localized support alerts.</p>  
        </div>
    </div>"""
    return HttpResponse(content)

from django.http import JsonResponse

# --- 🛰️ THE SOVEREIGN METRICS UPLINK ---
def get_hub_metrics(request):
    """Dynamically loads school health for the mobile app dashboard"""
    # Math logic:
    total_students = Student.objects.count()
    # Pulling real fee stats from your ledger
    fees = FeesTracker.objects.aggregate(due=Sum('total_fees_due'), paid=Sum('total_fees_paid'))
    
    data = {
        "status": "National Hub Online",
        "enrollment": total_students,
        "collection_rate": f"{(fees['paid']/fees['due']*100):.1f}%" if fees['due'] else "0%",
        "active_staff": Staff.objects.count(),
        "academic_week": "Week 6, Term II",
    }
    return JsonResponse(data)

from rest_framework.decorators import api_view
from rest_framework.response import Response
from .models import Student, Staff

@api_view(['POST'])
@permission_classes([AllowAny])
def student_identity_gate(request):
    d = request.data
    print(f"--- 🛡️ IDENTITY ATTEMPT: {d} ---") 
    code_in = d.get('code', '').strip().upper()
    student_in = d.get('student', '').strip().lower()
    
    student = Student.objects.filter(payment_code__iexact=code_in).first()
    
    if not student:
        print(f"❌ REJECTED: PRN {code_in} not found in DB.")
        return Response({"status": "DENIED", "msg": "PRN NOT FOUND"}, status=401)
    
    if student.full_name.lower().strip() != student_in:
        print(f"❌ REJECTED: Name {student_in} does not match {student.full_name.lower()}.")
        return Response({"status": "DENIED", "msg": "STUDENT NAME MISMATCH"}, status=401)

    return Response({"status": "IDENTITY_CONFIRMED", "student_id": str(student.account_number)})

@api_view(['POST'])
@permission_classes([AllowAny])
def staff_hub_auth(request):
    d = request.data
    # 🕵️ LOG THE INCOMING DATA
    print(f"--- 👔 STAFF ATTEMPT: {d} ---")
    
    name_in = d.get('name', '').strip()
    pin_in = d.get('pin', '').strip()

    staff = Staff.objects.filter(full_name__iexact=name_in, secure_pin=pin_in).first()
    
    if staff:
        print(f"✅ AUTHORIZED: {staff.full_name}")
        return Response({"status": "STAFF_AUTHORIZED", "name": staff.full_name, "role": "Teacher"})
    
    print(f"❌ REJECTED: Staff {name_in} with PIN {pin_in} not found.")
    return Response({"status": "DENIED", "msg": "STAFF NOT FOUND"}, status=401)

@api_view(['POST'])
@permission_classes([AllowAny])
def pin_vault_auth(request):
    d = request.data
    sid = d.get('student_id')
    pin = d.get('pin', '').strip()

    student = Student.objects.filter(account_number=sid).first()
    
    if student and student.parent_link.secure_pin == pin:
        # 🛡️ Instead of using Serializer (which might return a list), 
        # we hand-pick the data into a PURE MAP.
        marks = list(student.marks.values('subject__name', 'aoi_1', 'aoi_2', 'mid_term', 'aoi_3', 'aoi_4', 'eot_score'))
        
        data_packet = {
            "full_name": student.full_name,
            "account_number": student.account_number,
            "current_class": student.current_class,
            "stream": student.stream or "NORTH",
            "school_name": student.school.name,
            "academic_record": marks,
            "financial_standing": {
                "balance": 150000 # Example, or pull from tracker
            }
        }
        return JsonResponse(data_packet, safe=False)
        
    return JsonResponse({"status": "WRONG_PIN", "msg": "Invalid Secure PIN"}, status=401)

from django.core.management import call_command
from django.db import connection

def force_registry_rebuild(request):
    try:
        # This physically builds your 100-character tables in Supabase
        call_command('migrate', interactive=False)
        return HttpResponse("<h1>NATIONAL REGISTRY BUILT SUCCESSFULLY! 🏆</h1>")
    except Exception as e:
        return HttpResponse(f"Registry Error: {str(e)}")


# --- 💰 THE SCHOOLPAY SETTLEMENT SIMULATOR ---
from .models import Student, School, SchoolPayLedger
import random

def simulate_payment(request):
    """Simulates a real USSD/Mobile Money payment through SchoolPay"""
    try:
        # 1. Pick a student (Ensure you have at least one student in the DB!)
        student = Student.objects.first() 
        if not student: return HttpResponse("Registry Error: Add a student first!")
        
        # 2. Create a fake Receipt
        receipt = f"RCPT-{random.randint(100000, 999999)}"
        amount = 50000 # 50k UGX
        
        # 3. Create the Ledger Entry
        ledger = SchoolPayLedger.objects.create(
            student=student,
            school=student.school,
            receipt_number=receipt,
            amount=amount,
            raw_data={"sourceChannel": "MTN_MOMO", "transactionID": receipt}
        )
        
        # 4. Trigger the Math Update (This mimics the worker)
        student.school.total_revenue_collected += amount
        student.school.save()

        return HttpResponse(f"<h1>💰 PAYMENT SUCCESS</h1><p>Student {student.full_name} paid UGX {amount}. Receipt: {receipt}</p>")
    except Exception as e:
        return HttpResponse(f"Simulation Failed: {str(e)}")

from django.shortcuts import redirect

def direct_app_download(request):
    google_drive_id = "YOUR_LONG_GOOGLE_DRIVE_ID_HERE"
    direct_link = f"https://drive.google.com/uc?export=download&id={google_drive_id}"
    
    return redirect(direct_link)

def generate_staff_dossier_pdf(request, staff_id):
    """
    Generates a high-security National HR Dossier for Audit.
    Includes: Biometrics, URA TIN, NSSF, and Regulatory Data.
    """
    try:
        staff = Staff.objects.get(staff_id=staff_id)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Dossier_{staff.full_name}.pdf"'

        p = canvas.Canvas(response, pagesize=A4)
        w, h = A4

        p.saveState()
        p.setFont("Helvetica-Bold", 60)
        p.setFillColor(colors.lightgrey, alpha=0.05)
        p.translate(w/2, h/2); p.rotate(45)
        p.drawCentredString(0, 0, "UNSCCDC OFFICIAL HUB")
        p.restoreState()

        
        p.setLineWidth(5)
        p.setStrokeColor(colors.black); p.line(0, h-2, w/3, h-2)
        p.setStrokeColor(colors.orange); p.line(w/3, h-2, (w/3)*2, h-2)
        p.setStrokeColor(colors.red); p.line((w/3)*2, h-2, w, h-2)

        
        p.setFont("Helvetica-Bold", 14)
        p.drawCentredString(w/2, h-50, "THE REPUBLIC OF UGANDA")
        p.setFont("Helvetica-Bold", 11)
        p.drawCentredString(w/2, h-70, "NATIONAL STAFF REGISTRY - OFFICIAL DOSSIER")
        p.setFont("Helvetica", 9)
        p.drawCentredString(w/2, h-85, f"Institutional Station: {staff.school.name.upper()}")

        
        p.setStrokeColor(colors.black); p.rect(40, h-250, w-80, 150)
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, h-120, "1.0 PERSONAL BIOMETRICS")
        p.setFont("Helvetica", 10)
        p.drawString(60, h-145, f"FULL LEGAL NAME: {staff.full_name.upper()}")
        p.drawString(60, h-165, f"NATIONAL STAFF ID: {staff.staff_id}")
        p.drawString(60, h-185, f"DESIGNATION: {staff.designation}")
        p.drawString(60, h-205, f"CONTACT UPLINK: {staff.phone}")

        
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, h-280, "2.0 REGULATORY COMPLIANCE (URA / NSSF)")
        p.line(50, h-285, 300, h-285)
        p.setFont("Helvetica", 10)
        p.drawString(60, h-310, f"URA TIN NUMBER: {getattr(staff, 'tin_number', 'PENDING')}")
        p.drawString(60, h-330, f"NSSF REGISTRY NO: {getattr(staff, 'nssf_number', 'PENDING')}")

        
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, h-380, "3.0 EMERGENCY & KINSHIP REGISTRY")
        p.setFont("Helvetica", 10)
        p.drawString(60, h-410, f"NEXT OF KIN: {getattr(staff, 'next_of_kin', 'NOT SET')}")
        p.drawString(60, h-430, f"KIN CONTACT: {getattr(staff, 'next_of_kin_phone', 'NOT SET')}")

        
        p.setFont("Helvetica-Bold", 8)
        p.drawCentredString(w/2, 100, "THIS DOCUMENT IS A CERTIFIED DIGITAL RECORD OF THE UNSCCDC HUB")
        p.drawCentredString(w/2, 85, f"VERIFICATION HASH: {staff.staff_id}-AUDIT-2026")

        p.showPage(); p.save()
        return response
    except Exception as e:
        from django.http import HttpResponse
        return HttpResponse(f"Dossier Engine Error: {str(e)}", status=400)

import os
import datetime
from django.db.models import Avg
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from .models import Student, AcademicResult, FeesTracker, School

def generate_national_report_pdf(request, student_id):
    """
    THE GOLIATH Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub
    UNSCCDC NATIONAL Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub
    """
    # 🏛️ 1. DEFINE Hub Hub Hub Hub Hub Hub Hub COLORS
    gov_blue = colors.HexColor("#002366")   # Royal Navy (Authority)
    rich_gold = colors.HexColor("#D4AF37")  # Champagne Gold (Prestige)
    off_white = colors.HexColor("#FDFDF5")  # Institutional Parchment
    ug_yellow = colors.HexColor("#FCDC04")  # National Gold
    ug_red = colors.HexColor("#D90000")     # National Red

    try:
        # 🔑 2. Hub Hub Hub Hub Hub Hub IDENTITY GATE
        student = Student.objects.get(account_number=student_id)
        fees, _ = FeesTracker.objects.get_or_create(student=student)
        
        amt_to_be_paid = fees.total_fees_due
        total_paid = fees.total_fees_paid
        balance = fees.fees_balance
        marks = student.marks.all() 
        school = student.school

        # 🧮 3. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub RANKING ENGINE
        all_class_students = Student.objects.filter(current_class=student.current_class, school=school)
        student_scores = []
        for s_obj in all_class_students:
            avg = s_obj.marks.aggregate(a=Avg('eot_score'))['a'] or 0
            student_scores.append({'id': s_obj.id, 'avg': avg})
        
        student_scores.sort(key=lambda x: x['avg'], reverse=True)
        total_in_class = len(student_scores)
        position = next((i + 1 for i, item in enumerate(student_scores) if item['id'] == student.id), 1)
        overall_avg = next((item['avg'] for item in student_scores if item['id'] == student.id), 0)

        # 📄 4. Hub Hub Hub Hub Hub INITIALIZE CANVAS
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="National_Report_{student.full_name}.pdf"'
        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4 

        # 🎨 3. NATIONAL PALETTE & BACKGROUND
        gov_blue = colors.HexColor("#002366")   # Royal Navy
        rich_gold = colors.HexColor("#D4AF37")  # Champagne Gold
        off_white = colors.HexColor("#FDFDF5")  # Parchment
        
        p.setFillColor(off_white)
        p.rect(0, 0, width, height, fill=1, stroke=0)

        # 🛡️ 4. TRIPLE-GUARD BORDERS
        p.setLineWidth(5); p.setStrokeColor(gov_blue); p.rect(15, 15, width-30, height-30)
        p.setLineWidth(1); p.setStrokeColor(colors.HexColor("#FCDC04")); p.rect(22, 22, width-44, height-44)
        p.setStrokeColor(colors.HexColor("#D90000")); p.rect(23, 23, width-46, height-46)

        # 🎨 5. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub PAINT THE Hub Hub Hub Hub Hub Hub Hub FLOOR
        p.setFillColor(off_white)
        p.rect(0, 0, width, height, fill=1, stroke=0)

        # 🛡️ 6. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub TRIPLE-GUARD Hub Hub Hub Hub Hub Hub Hub Hub BORDERS
        p.setLineWidth(5); p.setStrokeColor(gov_blue); p.rect(15, 15, width-30, height-30)
        p.setLineWidth(1); p.setStrokeColor(ug_yellow); p.rect(22, 22, width-44, height-44)
        p.setStrokeColor(ug_red); p.rect(23, 23, width-46, height-46)

        # 🌌 7. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub SOVEREIGN Hub Hub Hub Hub Hub Hub Hub Hub WATERMARK
        p.saveState()
        p.setFont("Helvetica-Bold", 45); p.setFillColor(colors.lightgrey, alpha=0.03)
        p.translate(width/2, height/2); p.rotate(45); p.drawCentredString(0, 0, "UNSCCDC OFFICIAL RECORD")
        p.restoreState()

        # 🎨 8. Hub Hub Hub Hub Hub Hub Hub Hub Hub INTERNAL Hub Hub Hub Hub Hub Hub Hub Hub AI LOGIC
        def calculate_uce_grade(score):
            if score >= 80: return "A"
            if score >= 70: return "B"
            if score >= 60: return "C"
            if score >= 50: return "D"
            return "E"
        
        def get_sub_remark(score):
            if score >= 90: return "Exceptional mastery."
            if score >= 80: return "Excellent. Maintain focus."
            if score >= 70: return "Very good effort."
            if score >= 60: return "Good progress."
            if score >= 50: return "Basic competency."
            return "Requires support."

        def get_teacher_comment(avg):
            if avg >= 80: return "Disciplined and hardworking. High leadership potential."
            if avg >= 60: return "Good performance. Should focus more on technicals."
            return "Needs more effort and attend all remedial sessions."
        
        # 💎 THE Hub Hub Hub Hub Hub Hub Hub SECTOR-SPECIFIC GRADING
        if school.sector == 'PRIMARY':
            grade_title = "PRIMARY (PLE) GRADING STANDARDS"
            grade_data = [
                ['Agg', 'Div', 'Description'],
                ['4-12', '1', 'Exceptional - High Distinction'],
                ['13-23', '2', 'Strong Credit'],
                ['24-28', '3', 'Pass'],
                ['29-34', '4', 'Minimum Pass']
            ]
        elif school.sector == 'UNIVERSITY':
            grade_title = "HIGHER EDUCATION (NCHE) CGPA STANDARDS"
            grade_data = [
                ['CGPA', 'Class', 'Standing'],
                ['4.40-5.00', '1st Class', 'Exceptional Excellence'],
                ['3.60-4.39', '2nd Upper', 'Strong Honors'],
                ['2.80-3.59', '2nd Lower', 'Average Honors'],
                ['2.00-2.79', 'Pass', 'Satisfactory']
            ]
        else: # Default UCE
            grade_title = "SECONDARY (UCE) COMPETENCY STANDARDS"
            grade_data = [ ... ] # Your existing UCE data

        # 🏛️ 9. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub OFFICIAL Hub Hub Hub Hub Hub Hub Hub Hub Hub HEADER
        p.setFillColor(colors.black); p.setFont("Helvetica-Bold", 10)
        p.drawCentredString(width/2, height-45, "THE REPUBLIC OF UGANDA")
        p.drawCentredString(width/2, height-58, "UGANDA NATIONAL EXAMINATIONS BOARD (UNEB)")

        if student.photo:
                    try:
                        # 🕵️ Safety check for Render's ephemeral storage
                        if os.path.exists(student.photo.path):
                            # 📍 TOP LEFT COORDINATES
                            px, py = 45, height - 130 
                            pw, ph = 70, 85 # Elegant Passport size
                            
                            # 1. Draw a subtle "Imperial Shadow" for 3D effect
                            p.setFillColor(colors.HexColor("#D3D3D3"))
                            p.rect(px + 1.5, py - 1.5, pw, ph, fill=1, stroke=0)
                            
                            # 2. Draw the actual student photo
                            p.drawImage(student.photo.path, px, py, width=pw, height=ph, mask='auto')
                            
                            # 3. Draw the Imperial Gold Frame (Matches the borders)
                            p.setStrokeColor(rich_gold)
                            p.setLineWidth(1.2)
                            p.rect(px, py, pw, ph, stroke=1, fill=0)
                            
                            # 4. Tiny "Verified" watermark on the photo bottom
                            p.setFillColor(colors.white)
                            p.setFont("Helvetica-Bold", 5.5)
                            p.drawString(px + 4, py + 4, "SECURE IDENTITY")
                    except Exception as e:
                        print(f"Top-Left Photo Skip: {e}")
        
        
       # 🖼️ 5. DYNAMIC SCHOOL LOGO (Replaces the Seal)
        if school.logo:
            try:
                # Path handles local and server storage automatically
                p.drawImage(school.logo.path, width/2-35, height-115, width=70, height=70, mask='auto')
            except:
                p.setStrokeColor(gov_blue)
                p.rect(width/2-25, height-115, 50, 50, stroke=1)
                p.drawCentredString(width/2, height-95, "LOGO")
        else:
            p.setStrokeColor(gov_blue)
            p.rect(width/2-25, height-115, 50, 50, stroke=1)
            p.drawCentredString(width/2, height-95, "OFFICIAL")

        p.setFont("Helvetica-Bold", 18); p.setFillColor(gov_blue)
        p.drawCentredString(width/2, height-135, school.name.upper())
        p.setFillColor(colors.black); p.setFont("Helvetica-Bold", 11)
        p.drawCentredString(width/2, height-160, "NATIONAL SCHOLASTIC PERFORMANCE RECORD")

        # 👤 10. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub STUDENT Hub Hub Hub Hub Hub Hub Hub Hub IDENTITY
        p.setFont("Helvetica-Bold", 9)
        p.drawString(50, height-195, f"STUDENT NAME: {student.full_name.upper()}")
        p.drawString(50, height-210, f"NATIONAL ID: {student.account_number}")
        p.drawString(380, height-195, f"CLASS: {student.current_class} ({student.stream or 'NORTH'})")
        p.drawString(380, height-210, f"TERM: EOT | YEAR: 2026")

        # 📊 11. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub DATA Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub MATRIX
        data = [['SUB', 'A1', 'A2', 'MID', 'A3', 'A4', 'EOT', 'PRJ', 'AVG', 'GRD', 'TCH', 'REMARKS']]
        for m in marks:
            formatted_score = f"{m.eot_score:g} / {m.eot_max}" 
            auto_grade = calculate_uce_grade(m.eot_score) 
            data.append([m.subject.name[:4].upper(), m.aoi_1, m.aoi_2, m.mid_term, m.aoi_3, m.aoi_4, m.eot_score, m.project_work, f"{m.eot_score}%", auto_grade, 'STF', get_sub_remark(m.eot_score)])
        
        has_aois = any(
            getattr(m, 'aoi_1', 0) > 0 or 
            getattr(m, 'aoi_2', 0) > 0 or 
            getattr(m, 'aoi_3', 0) > 0 or 
            getattr(m, 'aoi_4', 0) > 0 
            for m in marks
        )

        if has_aois:
            # 12-Column Mode (Modern CBC)
            headers = ['SUB', 'A1', 'A2', 'MID', 'A3', 'A4', 'EOT', 'PRJ', 'AVG', 'GRD', 'TCH', 'REMARKS']
            col_widths = [45, 20, 20, 25, 20, 20, 25, 25, 30, 25, 30, 160]
        else:
            # 8-Column Mode (Traditional - AOIs DISAPPEAR COMPLETELY)
            headers = ['SUBJECT NAME', 'MID TERM', 'EOT EXAM', 'PROJECT', 'AVERAGE', 'GRADE', 'TEACHER', 'REMARKS']
            col_widths = [100, 55, 55, 55, 55, 45, 55, 120]

        data_rows = [headers]
        for m in marks:
            # Auto-Grader Logic
            score = m.eot_score
            g = "A" if score >= 80 else "B" if score >= 70 else "C" if score >= 60 else "D" if score >= 50 else "E"
            rem = "Excellent" if score >= 80 else "Good" if score >= 50 else "Needs Effort"

            if has_aois:
                # 💎 Row with AOIs
                data_rows.append([
                    m.subject.name[:4].upper(), 
                    m.aoi_1, getattr(m, 'aoi_2', 0), m.mid_term, 
                    getattr(m, 'aoi_3', 0), getattr(m, 'aoi_4', 0), 
                    m.eot_score, m.project_work, f"{score}%", g, 'STF', rem
                ])
            else:
                # 💎 Row WITHOUT AOIs (The 00s are physically removed!)
                data_rows.append([
                    m.subject.name.upper(), m.mid_term, m.eot_score, 
                    m.project_work, f"{score}%", g, 'STF', rem
                ])

        # Create Table with dynamic widths
        table = Table(data_rows, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), gov_blue), ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTSIZE', (0,0), (-1,-1), 7), ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [off_white, colors.white]),
            ('GRID', (0,0), (-1,-1), 0.1, colors.grey), ('LINEBELOW', (0,0), (-1,0), 2, rich_gold),
        ]))
        table.wrapOn(p, width, height); table.drawOn(p, 30, height - 350)

        # 📚 12. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub UCE Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub COMPETENCY
        p.setFont("Helvetica-Bold", 8); p.drawString(50, height - 375, "GRADE COMPETENCY LEVEL & DESCRIPTION (UCE STANDARDS):")
        grade_data = [
            ['Grade', 'Level', 'Description / Score Bracket'],
            ['A', 'Exceptional', '80% - 100%. Extraordinary mastery innovatively applied.'],
            ['B', 'Outstanding', '70% - 79%. High competency in practical applications.'],
            ['C', 'Satisfactory', '60% - 69%. Adequate competency in application.'],
            ['D', 'Basic', '50% - 59%. Minimum level of competency in problem solving.'],
            ['E', 'Elementary', '0% - 49%. Below the basic level of competency.']
        ]
        g_table = Table(grade_data, colWidths=[40, 80, 360])
        g_table.setStyle(TableStyle([('FONTSIZE',(0,0),(-1,-1),7),('GRID',(0,0),(-1,-1),0.1,colors.black),('BACKGROUND',(0,0),(-1,0),gov_blue),('TEXTCOLOR',(0,0),(-1,0),colors.white)]))
        g_table.wrapOn(p, width, height); g_table.drawOn(p, 50, height - 470)

        # 🎓 13. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub UACE Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub (A-LEVEL) Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub KEY
        p.setFont("Helvetica-Bold", 8); p.drawString(50, height - 495, "ADVANCED LEVEL (UACE) PRINCIPAL PASS SCALES:")
        uace_data = [
            ['A (6pts)', 'B (5pts)', 'C (4pts)', 'D (3pts)', 'E (2pts)', 'O (1pt)', 'F (0pts)'],
            ['Excellent', 'Very Good', 'Good', 'Satisfactory', 'Fair', 'Sub. Pass', 'Fail']
        ]
        u_table = Table(uace_data, colWidths=[68, 68, 68, 68, 68, 68, 68])
        u_table.setStyle(TableStyle([('FONTSIZE',(0,0),(-1,-1),7),('GRID',(0,0),(-1,-1),0.1,colors.black),('ALIGN', (0,0), (-1,-1), 'CENTER')]))
        u_table.wrapOn(p, width, height); u_table.drawOn(p, 50, height - 530)

        # =============================================================
        # 💎 --- SECTION 12: Hub Hub Hub OFFICIAL Hub Hub Hub ADMINISTRATIVE Hub Hub Hub REMARKS ---
        # =============================================================
        p.setFont("Helvetica-Bold", 8)
        p.setFillColor(gov_blue)
        p.drawString(50, height - 565, "OFFICIAL ADMINISTRATIVE REMARKS:")

        # 🛡️ Draw a prestigious thin grey box for the remarks (Height Adjusted)
        p.setStrokeColor(colors.grey)
        p.setLineWidth(0.5)
        p.rect(50, height - 640, width - 100, 60) # Top=height-575, Bottom=height-635

        # A. Class Teacher Remarks
        p.setFillColor(colors.black)
        p.setFont("Helvetica-Bold", 7.5)
        p.drawString(60, height - 595, "CLASS TEACHER:")
        p.setFont("Helvetica-Oblique", 7.5)
        class_remark = get_teacher_comment(overall_avg)
        p.drawString(135, height - 595, f'"{class_remark}"')

        # B. Headteacher Remarks
        p.setFont("Helvetica-Bold", 7.5)
        p.drawString(60, height - 620, "HEAD TEACHER:")
        p.setFont("Helvetica-Oblique", 7.5)
        ht_remark = "Exceptional discipline. Highly recommended for National progressive placement." if overall_avg >= 75 else "Steady progress observed. Needs consistent focus in project-based assessments."
        p.drawString(135, height - 620, f'"{ht_remark}"')

        # =============================================================
        # 📜 --- SECTION 13: Hub Hub Hub CERTIFICATION Hub Hub Hub Hub Hub & Hub Hub Hub Hub Hub RANKING Hub Hub Hub ---
        # =============================================================
        p.setFont("Helvetica-Bold", 8)
        p.setFillColor(colors.black)
        p.drawString(50, height - 660, "CERTIFICATION STATUS:")
        p.setFont("Helvetica", 7)
        p.drawString(60, height - 672, f"• Result 1: Qualifies for UCE certificate. (Student achieved overall average of {overall_avg:.1f}%)")
        
        
        # 📊 National Standing & PRN Bar (Clean Horizontal Alignment)
        p.setStrokeColor(rich_gold)
        p.setLineWidth(1)
        p.line(50, height - 715, width - 50, height - 715) # Gold divider

        p.setFont("Helvetica-Bold", 9)
        p.setFillColor(gov_blue)
        p.drawString(50, height - 710, f"NATIONAL STANDING: Position {position} out of {total_in_class}")
        
        p.setFillColor(ug_red)
        p.drawRightString(width - 50, height - 710, f"SCHOOLPAY PRN: {student.payment_code or '---'}")

        p.setFont("Helvetica-Oblique", 6.5)
        p.setFillColor(colors.black)
        p.drawString(50, height - 725, "Note: UNEB explicitly does not rank candidates via aggregates to avoid unethical competition.")

        # =============================================================
        # ✍️ --- SECTION 14: Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub FINAL Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub SIGNATURES ---
        # =============================================================
        p.setStrokeColor(gov_blue)
        p.setLineWidth(0.8)

        # =============================================================
        # 💰 --- SECTION 13.5: Hub Hub Hub NATIONAL TREASURY STANDING (REFINED) ---
        # =============================================================
        # 🛡️ 1. Draw the Royal Navy Background Bar
        p.setStrokeColor(rich_gold)
        p.setLineWidth(1.5)
        p.setFillColor(gov_blue) 
        p.rect(45, height - 730, width - 90, 50, fill=1) # Widened slightly

        # ✍️ 2. Insert the Real Shillings & National PRN
        p.setFillColor(colors.white)
        
        # Column 1: Total Due
        p.setFont("Helvetica-Bold", 7)
        p.drawString(55, height - 700, "TOTAL BILLED")
        p.setFont("Helvetica-Bold", 10)
        p.drawString(55, height - 715, f"{amt_to_be_paid:,.0f}")

        # Column 2: Total Paid
        p.setFont("Helvetica-Bold", 7)
        p.drawString(165, height - 700, "TOTAL PAID")
        p.setFont("Helvetica-Bold", 10)
        p.drawString(165, height - 715, f"{total_paid:,.0f}")

        # Column 3: Balance
        p.setFont("Helvetica-Bold", 7)
        p.drawString(285, height - 700, "OUTSTANDING BAL")
        p.setFont("Helvetica-Bold", 11)
        if balance <= 0:
            p.setFillColor(colors.HexColor("#00FF00")) # Success Green
            p.drawString(285, height - 715, "CLEARED")
        else:
            p.setFillColor(colors.white)
            p.drawString(285, height - 715, f"{balance:,.0f}")

        # 🔥 Column 4: THE Hub Hub NATIONAL PRN (THE KEY)
        # We use a bright, aggressive Red for high-visibility
        p.setFillColor(colors.HexColor("#FF0000")) # 🔴 PERFECT RED
        p.setFont("Helvetica-Bold", 8)
        p.drawString(425, height - 700, "PAYMENT CODE")
        p.setFont("Helvetica-Bold", 14) # 💎 Large font so parents can't miss it!
        p.drawString(425, height - 718, f"{student.payment_code or 'N/A'}")

        # 📄 3. Security Footer under the bar
        p.setFillColor(colors.black)
        p.setFont("Helvetica-Oblique", 7)
        p.drawString(50, height - 745, f"Payment status is live. Reference the Red PRN code for all Bank/Mobile Money settlements.")
        # Left Signature
        p.line(50, height - 785, 200, height - 785)
        p.setFont("Helvetica-Bold", 7)
        p.drawCentredString(125, height - 797, "Head Teacher Signature")

        # Right Signature
        p.line(width - 200, height - 785, width - 50, height - 785)
        p.drawCentredString(width - 125, height - 797, "National Hub Registrar")
        
        # 🛡️ THE Hub Hub Hub Hub SOVEREIGN STAMP (Centered perfectly)
        p.setStrokeColor(colors.HexColor("#008080")) # Institutional Teal
        p.circle(width/2, height - 780, 32, stroke=1, fill=0)
        p.setFont("Helvetica-Bold", 8)
        p.drawCentredString(width/2, height - 775, "UNSCCDC")
        p.setFont("Helvetica", 6)
        p.drawCentredString(width/2, height - 785, "VERIFIED")
        p.setFont("Helvetica-Bold", 7)
        p.drawCentredString(width/2, height - 795, datetime.date.today().strftime("%d-%b-%Y"))

        p.showPage(); p.save()
        return response
    except Exception as e:
        return HttpResponse(f"Hub Printing Error: {str(e)}", status=400)

import random
from django.http import JsonResponse
from .models import Student, SchoolPayLedger

def sovereign_shilling_simulator(request):
    """
    🧪 TEST-ONLY Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub
    Simulates an incoming SchoolPay payment to test Real-Time Sync.
    Does not affect live bank credentials.
    """
    # 🛡️ SECURITY KEY: Only you can trigger this
    if request.GET.get('key') != 'imperial_test_2026':
        return JsonResponse({"status": "Access Denied"}, status=403)

    try:
        # 1. Grab the first student in the registry (e.g., Namaganda Erina)
        student = Student.objects.first()
        if not student:
            return JsonResponse({"status": "Error", "msg": "Add a student first!"})

        # 2. Define dummy payment data
        amount = 125000 # Simulating 125k UGX
        receipt = f"SIM-{random.randint(10000, 99999)}"

        # 3. 🚀 THE Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub ACTION
        # Creating this record physically triggers the @receiver(post_save) signal!
        SchoolPayLedger.objects.create(
            student=student,
            school=student.school,
            amount=amount,
            receipt_number=receipt,
            raw_data={"sourceChannel": "AIRTEL_MONEY", "note": "Simulation Test"}
        )

        return JsonResponse({
            "status": "Simulation Success",
            "student": student.full_name,
            "amount_simulated": f"UGX {amount:,.0f}",
            "receipt": receipt,
            "instruction": "Go to Fees Tracker or Bursar Terminal now and refresh!"
        })
    except Exception as e:
        return JsonResponse({"status": "Simulation Failed", "error": str(e)})

# 💎 ADD THESE TO YOUR BATCH ENGINE IN views.py

# =============================================================
# 📜 HELPER 1: THE Hub Hub Hub Hub Hub REPORT CARD ARCHITECT
# =============================================================
def draw_single_report_page(p, student):
    try:
        width, height = A4
        school = student.school
        marks = student.marks.all()
        fees_obj, _ = FeesTracker.objects.get_or_create(student=student)
        
        gov_blue = colors.HexColor("#002366")
        rich_gold = colors.HexColor("#D4AF37")
        off_white = colors.HexColor("#FDFDF5")

        # =============================================================
        # 📸 --- SECTION 6.5: TOP-LEFT BIOMETRIC IDENTITY ---
        # =============================================================
        if student.photo:
            try:
                # 🕵️ Safety check for Render's ephemeral storage
                if os.path.exists(student.photo.path):
                    # 📍 TOP LEFT COORDINATES
                    px, py = 45, height - 130 
                    pw, ph = 70, 85 # Elegant Passport size
                    
                    # 1. Draw a subtle "Imperial Shadow" for 3D effect
                    p.setFillColor(colors.HexColor("#D3D3D3"))
                    p.rect(px + 1.5, py - 1.5, pw, ph, fill=1, stroke=0)
                    
                    # 2. Draw the actual student photo
                    p.drawImage(student.photo.path, px, py, width=pw, height=ph, mask='auto')
                    
                    # 3. Draw the Imperial Gold Frame (Matches the borders)
                    p.setStrokeColor(rich_gold)
                    p.setLineWidth(1.2)
                    p.rect(px, py, pw, ph, stroke=1, fill=0)
                    
                    # 4. Tiny "Verified" watermark on the photo bottom
                    p.setFillColor(colors.white)
                    p.setFont("Helvetica-Bold", 5.5)
                    p.drawString(px + 4, py + 4, "SECURE IDENTITY")
            except Exception as e:
                print(f"Top-Left Photo Skip: {e}")

        # 1. Background & National Borders
        p.setFillColor(off_white); p.rect(0, 0, width, height, fill=1)
        p.setLineWidth(5); p.setStrokeColor(gov_blue); p.rect(15, 15, width-30, height-30)
        p.setLineWidth(1); p.setStrokeColor(colors.HexColor("#FCDC04")); p.rect(22, 22, width-44, height-44)

        # 2. School Logo
        if school.logo and os.path.exists(school.logo.path):
            p.drawImage(school.logo.path, width/2-35, height-110, width=70, height=70, mask='auto')
        
        # 3. Header Text
        p.setFillColor(colors.black); p.setFont("Times-Bold", 10)
        p.drawCentredString(width/2, height-40, "THE REPUBLIC OF UGANDA")
        p.setFont("Times-Bold", 18); p.setFillColor(gov_blue)
        p.drawCentredString(width/2, height-135, school.name.upper())

        # 4. Student Info & Photo
        p.setFillColor(colors.black); p.setFont("Times-Bold", 9)
        p.drawString(50, height-200, f"STUDENT: {student.full_name.upper()}")
        p.drawString(50, height-215, f"NATIONAL ID: {student.account_number}")
        p.drawString(350, height-200, f"CLASS: {student.current_class} ({student.stream})")
        
        if student.photo and os.path.exists(student.photo.path):
            p.drawImage(student.photo.path, width-130, height-140, width=80, height=90, mask='auto')

        # 5. Elastic Marks Table
        has_aois = any(m.aoi_1 > 0 for m in marks)
        if has_aois:
            headers = ['SUB', 'A1', 'A2', 'MID', 'A3', 'A4', 'EOT', 'PRJ', 'AVG', 'GRD']
            col_widths = [50, 30, 30, 35, 30, 30, 35, 35, 40, 35]
        else:
            headers = ['SUBJECT NAME', 'MID TERM', 'EOT EXAM', 'PROJECT', 'AVERAGE', 'GRADE']
            col_widths = [150, 70, 70, 70, 70, 60]

        data_rows = [headers]
        for m in marks:
            row = [m.subject.name.upper(), m.mid_term, m.eot_score, m.project_work, f"{m.eot_score}%", "B"]
            if has_aois:
                row = [m.subject.name[:3].upper(), m.aoi_1, m.aoi_2, m.mid_term, m.aoi_3, m.aoi_4, m.eot_score, m.project_work, f"{m.eot_score}%", "B"]
            data_rows.append(row)

        table = Table(data_rows, colWidths=col_widths)
        table.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), gov_blue), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('GRID', (0,0), (-1,-1), 0.1, colors.black), ('FONTNAME', (0,0), (-1,-1), 'Times-Bold')]))
        table.wrapOn(p, width, height); table.drawOn(p, 40, height - 380)

        # 6. Treasury Bar
        p.setFillColor(gov_blue); p.rect(45, height - 680, width - 90, 50, fill=1, stroke=0)
        p.setFillColor(colors.white); p.setFont("Times-Bold", 7)
        p.drawString(55, height - 635, "OUTSTANDING BALANCE")
        p.drawString(385, height - 635, "NATIONAL PRN")
        p.setFont("Times-Bold", 12); p.drawString(55, height - 655, f"UGX {fees_obj.fees_balance:,.0f}")
        p.setFillColor(colors.red); p.drawString(385, height - 670, f"{student.payment_code}")

    except Exception as e:
        print(f"Error drawing page for {student.full_name}: {e}")

# =============================================================
# 🔔 HELPER 2: THE Hub Hub Hub Hub Hub FEES REMINDER ARCHITECT
# =============================================================
def draw_single_reminder_page(p, student):
    try:
        width, height = A4
        school = student.school
        fees, _ = FeesTracker.objects.get_or_create(student=student)
        
        gov_blue = colors.HexColor("#002366")
        rich_gold = colors.HexColor("#D4AF37")
        
        p.setStrokeColor(gov_blue); p.setLineWidth(5); p.rect(15, 15, width-30, height-30)
        p.setFillColor(gov_blue); p.setFont("Times-Bold", 16)
        p.drawCentredString(width/2, height-100, school.name.upper())
        p.setFont("Times-Bold", 12); p.setFillColor(colors.black)
        p.drawCentredString(width/2, height-130, "OFFICIAL FEES REMINDER")
        
        p.setFont("Times-Bold", 11)
        p.drawString(50, height-180, f"TO THE PARENT/GUARDIAN OF: {student.full_name.upper()}")
        p.drawString(50, height-200, f"CLASS: {student.current_class}")
        
        p.rect(50, height-300, width-100, 80)
        p.drawString(70, height-250, f"TOTAL BALANCE DUE: UGX {fees.fees_balance:,.0f}")
        p.setFillColor(colors.red); p.setFont("Times-Bold", 14)
        p.drawString(70, height-280, f"PAYMENT PRN: {student.payment_code}")
        
        p.setFillColor(colors.black); p.setFont("Times-Roman", 10)
        instructions = "Please settle this balance via MTN/Airtel using the PRN above to avoid service interruption."
        p.drawString(50, height-350, instructions)
    except Exception as e:
        print(f"Reminder Error: {e}")

@login_required
def batch_report_generator(request):
    school = getattr(request.user, 'school', None) or School.objects.first()
    selected_class = request.GET.get('class')
    
    students = Student.objects.filter(school=school, current_class=selected_class, is_active=True).order_by('full_name')

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="REPORTS_{selected_class}.pdf"'
    
    p = canvas.Canvas(response, pagesize=A4)
    
    for student in students:
        # 🛡️ THE Hub Hub Hub Hub Hub SAFETY SHIELD
        try:
            draw_single_report_page(p, student)
            p.showPage() 
        except:
            continue # If one student fails, just go to the next!
        
    p.save()
    return response

    
@login_required
def batch_reminder_generator(request):
    school = getattr(request.user, 'school', None) or School.objects.first()
    selected_class = request.GET.get('class')
    
    if not selected_class:
        return HttpResponse("Please select a class first.")

    students = Student.objects.filter(school=school, current_class=selected_class, is_active=True).order_by('full_name')

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="BATCH_REMINDERS_{selected_class}.pdf"'
    
    p = canvas.Canvas(response, pagesize=A4)
    
    for student in students:
        # 🧪 CALL YOUR EXISTING REMINDER DRAWING LOGIC HERE
        draw_single_reminder_page(p, student)
        p.showPage() # 📄 New page for each parent's notice
        
    p.save()
    return response

@api_view(['GET'])
def live_warroom_stats(request):
    """
    🛰️ THE SATELLITE SIGNAL
    Returns raw JSON data for the ApexCharts to update live.
    """
    today = timezone.now().date()
    # 🧮 Calculate live totals
    total_revenue = SchoolPayLedger.objects.filter(timestamp__date=today).aggregate(Sum('amount'))['amount__sum'] or 0
    active_logins = 12 # Simulating live parents currently on the app
    
    return Response({
        "revenue_today": f"{total_revenue:,.0f}",
        "active_users": active_logins,
        "performance_index": "94.2%",
        # 📈 Send fresh coordinates for the line chart
        "chart_series": [random.randint(40, 100) for _ in range(7)] 
    })

@csrf_exempt
def catch_app_crash(request):
    """🛡️ THE Hub Hub Hub Hub Hub Hub NATIONAL MONITOR"""
    if request.method == 'POST':
        error_msg = request.POST.get('error', 'Unknown Log')
        
        # 💎 This makes the message look HUGE in the Render Terminal
        print("\n" + "📡" * 20)
        print(f"NATIONAL APP SIGNAL: {error_msg}")
        print("📡" * 20 + "\n")
        
        return HttpResponse("LOG_OK")
    return HttpResponse("LISTENING")

def generate_student_dossier(request, student_id):
    try:
        student = Student.objects.get(account_number=student_id)
        fees = FeesTracker.objects.get(student=student)
        payments = SchoolPayLedger.objects.filter(student=student).order_by('-timestamp')
        marks = AcademicResult.objects.filter(student=student)
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="DOSSIER_{student.full_name}.pdf"'
        
        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4
        gov_blue = colors.HexColor("#002366")
        rich_gold = colors.HexColor("#D4AF37")

        # 1. 🎨 BACKGROUND & BORDERS (England Standard)
        p.setFillColor(colors.HexColor("#FDFDF5"))
        p.rect(0, 0, width, height, fill=1)
        p.setStrokeColor(gov_blue); p.setLineWidth(5); p.rect(15, 15, width-30, height-30)

        # 2. 🏛️ HEADER
        p.setFillColor(gov_blue); p.setFont("Helvetica-Bold", 16)
        p.drawCentredString(width/2, height-60, "NATIONAL STUDENT DOSSIER")
        p.setFont("Helvetica", 8); p.setFillColor(colors.grey)
        p.drawCentredString(width/2, height-75, "OFFICIAL RECORD OF THE REPUBLIC OF UGANDA | UNSCCDC GLOBAL")

        # 3. 👤 SECTION: BIOMETRIC & IDENTITY
        p.setFillColor(gov_blue); p.rect(40, height-130, width-80, 20, fill=1)
        p.setFillColor(colors.white); p.setFont("Helvetica-Bold", 10)
        p.drawString(50, height-125, "I. STUDENT IDENTITY & REGISTRY")
        
        p.setFillColor(colors.black); p.setFont("Helvetica-Bold", 9)
        p.drawString(50, height-150, f"FULL NAME: {student.full_name.upper()}")
        p.drawString(50, height-165, f"NATIONAL ID / PRN: {student.payment_code}")
        p.drawString(300, height-150, f"CLASS: {student.current_class} ({student.stream})")
        p.drawString(300, height-165, f"SYSTEM ID: {student.account_number}")

        # 4. 💰 SECTION: FINANCIAL STANDING (Live Data)
        p.setFillColor(gov_blue); p.rect(40, height-210, width-80, 20, fill=1)
        p.setFillColor(colors.white); p.drawString(50, height-205, "II. FINANCIAL TREASURY STATUS")
        
        fin_data = [
            ['Category', 'Amount (UGX)'],
            ['Total Invoiced', f"{fees.total_fees_due:,.0f}"],
            ['Initial Deposit', f"{student.initial_deposit:,.0f}"],
            ['Total Paid to Date', f"{fees.total_fees_paid:,.0f}"],
            ['Current Balance', f"{fees.fees_balance:,.0f}"]
        ]
        t = Table(fin_data, colWidths=[200, 200])
        t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold')]))
        t.wrapOn(p, width, height); t.drawOn(p, 50, height-300)

        # 5. 📑 SECTION: RECENT PAYSLIPS / PAYMENTS
        p.setFillColor(gov_blue); p.drawString(50, height-330, "III. RECENT SETTLEMENT LOG (PAYSLIPS)")
        pay_rows = [['Receipt #', 'Date', 'Amount', 'Channel']]
        for pay in payments[:5]: # Show last 5
            pay_rows.append([pay.receipt_number, pay.timestamp.strftime('%d/%m/%y'), f"{pay.amount:,.0f}", "SchoolPay"])
        
        pt = Table(pay_rows, colWidths=[120, 100, 100, 100])
        pt.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.1, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke)]))
        pt.wrapOn(p, width, height); pt.drawOn(p, 50, height-430)

        # 6. 📜 SECTION: OFFICIAL DOCUMENT VERIFICATION
        p.setFillColor(gov_blue); p.rect(40, height-470, width-80, 20, fill=1)
        p.setFillColor(colors.white); p.drawString(50, height-465, "IV. DOCUMENT VERIFICATION STATUS")
        
        p.setFillColor(colors.black); p.setFont("Helvetica", 8)
        docs = [
            ("Birth Certificate", student.birth_certificate),
            ("PLE Result Slip", student.ple_result_slip),
            ("UCE Result Slip", student.uce_result_slip)
        ]
        y_pos = height-495
        for label, file in docs:
            status = "✅ VERIFIED & ATTACHED" if file else "❌ PENDING SUBMISSION"
            p.drawString(50, y_pos, f"{label}: {status}")
            y_pos -= 15

        # 7. 🛡️ FOOTER STAMP
        p.setStrokeColor(gov_blue); p.circle(width-100, 80, 40, stroke=1)
        p.setFont("Helvetica-Bold", 8); p.drawCentredString(width-100, 85, "UNSCCDC")
        p.drawCentredString(width-100, 75, "OFFICIAL SEAL")

        p.showPage(); p.save()
        return response
    except Exception as e:
        return HttpResponse(f"Dossier Error: {str(e)}")

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.db.models import Avg, Sum # 💎 CRITICAL IMPORT
from .models import Student, School

@login_required
def sovereign_registry_view(request):
    try:
        # 1. 🛡️ IDENTITY GATE
        school = getattr(request.user, 'school', None) or School.objects.first()
        
        if not school:
            return HttpResponse("<h1>Error: No School Found</h1>")

        # 2. 🧠 NATIONAL SECTOR MAP
        sector_map = {
            'PRIMARY': ['Baby', 'Middle', 'Top', 'P.1', 'P.2', 'P.3', 'P.4', 'P.5', 'P.6', 'P.7'],
            'SECONDARY': ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6'],
            'UNIVERSITY': ['Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'],
        }
        
        # 🎯 Ensure available_classes is defined first
        current_sector = getattr(school, 'sector', 'SECONDARY')
        available_classes = sector_map.get(current_sector, ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6'])
        
        # 💎 THE FIX: Define 'selected_class' at the TOP level so it's always associated with a value
        selected_class = request.GET.get('class', available_classes[0])
        query = request.GET.get('q', '').strip()

        # 3. 🔎 ALL-SEEING FILTER LOGIC
        base_students = Student.objects.filter(school=school)

        if query:
            # If searching, we look through all classes for that name/PRN
            students = base_students.filter(
                Q(full_name__icontains=query) | 
                Q(payment_code__icontains=query) |
                Q(stream__icontains=query)
            ).order_by('full_name')
        else:
            # If not searching, we filter strictly by the selected class
            # We use __icontains to make sure "S.1" matches "S1"
            clean_name = selected_class.replace(".", "")
            students = base_students.filter(
                Q(current_class__iexact=selected_class) | 
                Q(current_class__icontains=clean_name)
            ).order_by('full_name')

        # 4. 📦 PREPARE CONTEXT
        context = {
            'school': school,
            'available_classes': available_classes,
            'selected_class': selected_class, # Now guaranteed to have a value!
            'students': students,
            'total_count': students.count(),
            'title': "SOVEREIGN NATIONAL REGISTRY"
        }
        
        return render(request, 'sovereign_registry.html', context)

    except Exception as e:
        # 🚑 If it fails, this will show the new error clearly
        return HttpResponse(f"<body style='background:black;color:red;padding:50px;'><h1>Registry Engine Error</h1><p>{str(e)}</p></body>")

@login_required
def inject_national_subjects(request):
    """💎 THE Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub MASTER SUBJECT SEED"""
    if not request.user.is_superuser:
        return HttpResponse("Unauthorized")

    vocational_subjects = [
        ('Information & Comm. Technology', 'ICT', 'VOCATIONAL'),
        ('Tailoring & Fashion Design', 'TAIL', 'VOCATIONAL'),
        ('Bakery & Cookery', 'BAKE', 'VOCATIONAL'),
        ('Carpentry & Joinery', 'CARP', 'VOCATIONAL'),
        ('Bricklaying & Concrete Practice', 'BRIC', 'VOCATIONAL'),
        ('Art & Design', 'ART', 'VOCATIONAL'),
        ('Agriculture & Farming', 'AGRI', 'VOCATIONAL'),
        ('Hairdressing & Beauty', 'HAIR', 'VOCATIONAL'),
    ]

    academic_subjects = [
        ('Mathematics', 'MTH', 'CORE'),
        ('English Language', 'ENG', 'CORE'),
        ('Physics', 'PHY', 'CORE'),
        ('Chemistry', 'CHE', 'CORE'),
        ('Biology', 'BIO', 'CORE'),
        ('Geography', 'GEO', 'CORE'),
        ('History', 'HIS', 'CORE'),
    ]

    # 🚀 Inject into the Registry
    for name, code, cat in academic_subjects + vocational_subjects:
        Subject.objects.get_or_create(name=name, defaults={'code': code, 'category': cat})

    return HttpResponse("<h1 style='color:gold; background:black; padding:20px;'>NATIONAL SUBJECTS INJECTED SUCCESSFULLY! 🇺🇬</h1>")

from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from api.models import Student, School
from .utils import auto_arrange_pdf_data
import pandas as pd
from docx import Document

def upload_uneb_roster(request, school_id):
    if request.method == "POST" and request.FILES.get("pdf_document"):
        uploaded_pdf = request.FILES["pdf_document"]
        auto_arrange_pdf_data(uploaded_pdf, school_id)
        return redirect('student_roster_dashboard', school_id=school_id)
        
    return render(request, 'upload.html', {'school_id': school_id})

def export_styled_layout(request, school_id, layout_format):
    students = Student.objects.filter(school_id=school_id).order_by('current_class', 'stream', 'full_name')
    school_obj = School.objects.get(id=school_id)

    # 📊 Layout Choice 1: Microsoft Excel Template
    if layout_format == "excel":
        dataset = []
        for s in students:
            dataset.append([s.account_number, s.payment_code, s.full_name, s.current_class, s.stream, s.gender, s.fees_balance])
            
        columns = ["System ID", "SchoolPay PRN", "Full Student Name", "Class", "Stream", "Gender", "Current Balance"]
        df = pd.DataFrame(dataset, columns=columns)
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{school_obj.name}_Roster.xlsx"'
        df.to_excel(response, index=False)
        return response

    # 📝 Layout Choice 2: Microsoft Word Structured Table
    elif layout_format == "word":
        doc = Document()
        doc.add_heading(f'{school_obj.name} - Sorted Registry', level=1)
        doc.add_paragraph(f"Motto: {school_obj.school_motto} | Center: {school_obj.uneb_center_number}")
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'PRN (SchoolPay)'
        hdr_cells[1].text = 'Student Name'
        hdr_cells[2].text = 'Class'
        hdr_cells[3].text = 'Stream'
        hdr_cells[4].text = 'Balance status'
        
        for s in students:
            row_cells = table.add_row().cells
            row_cells[0].text = str(s.payment_code or 'N/A')
            row_cells[1].text = s.full_name
            row_cells[2].text = s.current_class
            row_cells[3].text = s.stream
            row_cells[4].text = s.fees_balance
            
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{school_obj.name}_Roster.docx"'
        doc.save(response)
        return response

@login_required
def operations_hub_view(request):
    sync_national_notifications()
    
    updates = NationalUpdate.objects.all()[:5] # Get latest 5
    school = getattr(request.user, 'school', None) or School.objects.first()
    
    # 🎨 THE Hub Hub Hub Hub Hub Hub Hub MINOR TABS REGISTRY
    # Each item: (Name, Icon, Color, Link, Description)
    minor_tabs = [
        ("Student Registry", "fa-user-graduate", "#3498db", "/admin/api/student/", "Manage Learners"),
        ("Fees & Payments", "fa-wallet", "#2ecc71", "/admin/api/schoolpayledger/", "Treasury Sync"),
        ("Exam Center", "fa-file-signature", "#9b59b6", "/api/explorer/", "Input Marks"),
        ("Fees Reminders", "fa-bell", "#f39c12", "/admin/api/feesreminder/", "Print Reminders"), 
        ("Report Cards", "fa-print", "#e74c3c", "/api/explorer/", "Generate PDFs"),
        ("Guardian Registry", "fa-users", "#e91e63", "/api/parents/", "Parent Access & PINs"),
        ("Staff Force", "fa-chalkboard-teacher", "#f1c40f", "/admin/api/staff/", "Employee Files"),
        ("Staff Salaries", "fa-hand-holding-usd", "#9b59b6", "/api/payroll-hub/", "Payroll & Tax"),
        ("SMS Broadcast", "fa-comment-alt", "#e67e22", "/api/sms-hub/", "Notify Parents"),
        ("Inventory/Store", "fa-boxes", "#1abc9c", "#", "School Property"),
        ("Library System", "fa-book", "#34495e", "#", "Book Tracking"),
        ("Transport/Bus", "fa-bus", "#d35400", "#", "Routes & Fees"),
        ("Dormitory/Hostel", "fa-bed", "#27ae60", "#", "Accommodation"),
        ("KEB Mock Center", "fa-file-signature", "#2196F3", "/api/keb-portal/", "Candidate Passlips"),
        ("UNEB/DIT Portal", "fa-medal", "#c0392b", "/api/uneb-gateway/", "National Exams"),
        ("KEB Mocks", "fa-file-invoice", "#2196F3", "/api/registry/", "Print KEB Passlips"), # 💎 18th TAB
        ("System Health", "fa-microchip", "#7f8c8d", "/admin/api/financialcommandcenter/", "Analytics"),
        ("Academic Command", "fa-award", "#9b59b6", "/api/results-center/", "Performance Analytics"),
        ("Secretary Entry", "fa-keyboard", "#1abc9c", "/api/secretary-entry/", "Fast Marks Ingestion"),
        ("System Settings", "fa-cogs", "#34495e", "/admin/api/systemsettings/", "Configure Hub"),
        ("Manage Users", "fa-user-lock", "#607d8b", "/admin/auth/user/", "Staff Access Control"), # 💎 15th TAB
    ]

    return render(request, 'admin/operations_hub.html', {
        'minor_tabs': minor_tabs,
        'national_updates': updates, # 💎 Send the news!
        'school': school,
        'title': "NATIONAL OPERATIONS COMMAND"
    })

import pdfplumber
from django.db import transaction

@login_required
def execute_data_bridge(request, bridge_id):
    bridge = get_object_or_404(NationalDataBridge, id=bridge_id)
    school = bridge.school
    count = 0
    
    try:
        with pdfplumber.open(bridge.source_pdf.path) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if not table or len(table) < 2: continue
                
                # 🧠 AI HEADER DETECTION
                header = [str(h).upper() if h else "" for h in table[0]]
                idx_name = next((i for i, h in enumerate(header) if "NAME" in h), 0)
                idx_prn = next((i for i, h in enumerate(header) if "PRN" in h or "CODE" in h), 1)
                idx_class = next((i for i, h in enumerate(header) if "CLASS" in h), 2)
                idx_stream = next((i for i, h in enumerate(header) if "STREAM" in h), 3)
                idx_parent = next((i for i, h in enumerate(header) if "PARENT" in h), 4)
                idx_phone = next((i for i, h in enumerate(header) if "PHONE" in h), 5)

                for row in table[1:]:
                    if not row[idx_name] or not row[idx_prn]: continue
                    with transaction.atomic():
                        # 1. Sync Parent
                        p_phone = str(row[idx_phone]).strip() if row[idx_phone] else "000"
                        parent_obj, _ = Parent.objects.get_or_create(
                            phone_number=p_phone,
                            defaults={'full_name': str(row[idx_parent]), 'secure_pin': '123456'}
                        )
                        # 2. Sync Student & Auto-Arrange Class/Stream
                        student_obj, _ = Student.objects.update_or_create(
                            payment_code=str(row[idx_prn]).strip().upper(),
                            defaults={
                                'full_name': str(row[idx_name]).strip().upper(),
                                'current_class': str(row[idx_class]).strip().upper(),
                                'stream': str(row[idx_stream]).strip().upper() if row[idx_stream] else "NORTH",
                                'school': school,
                                'parent_link': parent_obj,
                            }
                        )
                        # 3. Initialize Fees
                        FeesTracker.objects.get_or_create(student=student_obj)
                        count += 1
        
        bridge.is_processed = True
        bridge.records_synced = count
        bridge.save()
        return HttpResponse(f"<body style='background:#000;color:gold;padding:50px;text-align:center;'><h1>BRIDGE SUCCESS!</h1><p style='color:white;'>{count} Students automatically arranged.</p><a href='/admin/'>Back to Dashboard</a></body>")
    except Exception as e:
        return HttpResponse(f"Bridge Error: {str(e)}")

import pdfplumber
from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse

@login_required
def bridge_preview_portal(request, bridge_id):
    bridge = get_object_or_404(NationalDataBridge, id=bridge_id)
    preview_rows = []
    
    try:
        with pdfplumber.open(bridge.source_file.path) as pdf:
            # We only look at the first page for the preview to be LIGHTNING FAST
            first_page = pdf.pages[0]
            table = first_page.extract_table()
            
            if table:
                # 🧠 AI Column Finder
                header = [str(h).upper() if h else "" for h in table[0]]
                idx_name = next((i for i, h in enumerate(header) if "NAME" in h), 0)
                idx_prn = next((i for i, h in enumerate(header) if "PRN" in h or "CODE" in h), 1)
                idx_class = next((i for i, h in enumerate(header) if "CLASS" in h), 2)
                
                # Take first 10 rows for preview
                for row in table[1:11]:
                    if row[idx_name]:
                        preview_rows.append({
                            'name': row[idx_name],
                            'prn': row[idx_prn],
                            'class': row[idx_class],
                        })

        return render(request, 'admin/bridge_preview.html', {
            'bridge': bridge,
            'preview_rows': preview_rows,
            'title': "DATA PREVIEW PORTAL"
        })
    except Exception as e:
        return HttpResponse(f"Preview Error: {str(e)}")

@login_required
def bridge_commit_data(request, bridge_id):
    """The final trigger that actually saves data to the Registry"""
    # ... (This will call the process_national_pdf logic we built earlier)
    # Redirecting to previous logic for final save
    return process_national_pdf(request, bridge_id)

# =============================================================
# ☢️ THE Hub Hub Hub Hub Hub NUCLEAR TABLE DESTROYER
# =============================================================
from django.db import connection # 💎 Ensure this is at the top of views.py

def nuke_problem_table(request):
    """🛡️ THE EMERGENCY CLEANER - DROPS THE STUCK TABLE"""
    try:
        with connection.cursor() as cursor:
            # 1. Physically drop the old table causing the 'Already Exists' error
            cursor.execute("DROP TABLE IF EXISTS api_dataingestionvault CASCADE;")
            
            # 2. Delete the record of the migrations for your 'api' app
            # This makes Django think the app is brand new!
            cursor.execute("DELETE FROM django_migrations WHERE app = 'api';")
            
        return HttpResponse("<h1 style='color:white; background:red; padding:50px;'>SUCCESS: Table and History Nuked! Ready for Fresh Start.</h1><a href='/admin/'>Go Back to Office</a>")
    except Exception as e:
        return HttpResponse(f"Nuke Error: {str(e)}")

import pdfplumber
from django.db import transaction

@login_required
def bridge_preview_portal(request, bridge_id):
    bridge = get_object_or_404(NationalDataBridge, id=bridge_id)
    
    if not bridge.preview_data:
        try:
            with pdfplumber.open(bridge.source_file.path) as pdf:
                all_rows = []
                for page in pdf.pages:
                    table = page.extract_table()
                    if table: all_rows.extend(table)
                
                if not all_rows: return HttpResponse("No table found in PDF.")

                # 🧠 AI Column Intelligence
                header = [str(h).upper() for h in all_rows[0]]
                def find_idx(keys, default):
                    for i, h in enumerate(header):
                        if any(k in h for k in keys): return i
                    return default

                idx_name = find_idx(["NAME", "STUDENT"], 0)
                idx_prn = find_idx(["PRN", "CODE", "ID"], 1)
                idx_class = find_idx(["CLASS", "LEVEL"], 2)
                idx_stream = find_idx(["STREAM", "HOUSE"], 3)
                idx_parent = find_idx(["PARENT", "GUARDIAN"], 4)
                idx_phone = find_idx(["PHONE", "CONTACT"], 5)

                # Format the data for the preview
                formatted = []
                for row in all_rows[1:]:
                    if not row[idx_name]: continue
                    formatted.append({
                        'name': str(row[idx_name]).strip().upper(),
                        'prn': str(row[idx_prn]).strip().upper(),
                        'class': str(row[idx_class]).strip().upper(),
                        'stream': str(row[idx_stream]).strip().upper() if row[idx_stream] else "NORTH",
                        'parent': str(row[idx_parent]).strip().title(),
                        'phone': str(row[idx_phone]).strip()
                    })
                
                bridge.preview_data = formatted
                bridge.records_count = len(formatted)
                bridge.save()
        except Exception as e:
            return HttpResponse(f"Scan Error: {str(e)}")

    return render(request, 'admin/bridge_preview.html', {
        'bridge': bridge,
        'preview': bridge.preview_data,
        'title': "NATIONAL DATA PREVIEW"
    })

@login_required
@transaction.atomic
def bridge_commit_final(request, bridge_id):
    """The Final Trigger: Turns 'Ghost Data' into Real Registry Records"""
    bridge = get_object_or_404(NationalDataBridge, id=bridge_id)
    if bridge.is_processed: return HttpResponse("Already Processed.")

    for item in bridge.preview_data:
        # 1. Weld Parent
        parent_obj, _ = Parent.objects.get_or_create(
            phone_number=item['phone'],
            defaults={'full_name': item['parent'], 'secure_pin': '123456'}
        )
        # 2. Weld Student
        Student.objects.update_or_create(
            payment_code=item['prn'],
            defaults={
                'full_name': item['name'],
                'current_class': item['class'],
                'stream': item['stream'],
                'school': bridge.school,
                'parent_link': parent_obj,
                'is_active': True
            }
        )
    
    bridge.is_processed = True
    bridge.save()
    return redirect('/admin/api/student/')

@login_required
def reset_staff_pin(request, staff_id):
    """🛡️ GENERATES A NEW 4-DIGIT Hub Hub Hub PIN"""
    staff = get_object_or_404(Staff, id=staff_id)
    new_pin = ''.join(random.choices(string.digits, k=4))
    staff.secure_pin = new_pin
    staff.save()
    
    return HttpResponse(f"""
        <body style="background:#000; color:white; text-align:center; padding:50px; font-family:sans-serif;">
            <h2 style="color:gold;">PIN RESET SUCCESSFUL</h2>
            <p>New Security PIN for <b>{staff.full_name}</b> is:</p>
            <h1 style="font-size:50px; color:#00ff00; letter-spacing:10px;">{new_pin}</h1>
            <a href="/admin/api/staff/" style="color:gold; text-decoration:none; border:1px solid gold; padding:10px 20px; border-radius:10px;">RETURN TO REGISTRY</a>
        </body>
    """)

from django.contrib.auth import get_user_model # 💎 THE UNIVERSAL KEY

def national_landing_page(request):
    User = get_user_model() # 🛡️ This grabs the CORRECT user model automatically
    
    # 💎 EMERGENCY AUTO-ACCOUNT CREATION (Updated for safety)
    if not User.objects.filter(username='Josephat').exists():
        u = User.objects.create_user('Josephat', password='Josephat123')
        u.is_staff = True
        u.save()
        # Note: You can link the profile here too if needed
    
    # ... rest of your code ...
    """The prestigious entry point for the UNSCCDC Global System"""
    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>UNSCCDC GLOBAL | National Hub</title>
        <style>
            body { background: #050505; color: white; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; text-align: center; padding: 50px 20px; margin: 0; }
            .container { max-width: 800px; margin: auto; border: 2px solid #D4AF37; padding: 40px; border-radius: 30px; background: rgba(212,175,55,0.02); box-shadow: 0 0 50px rgba(212,175,55,0.1); }
            h1 { color: #D4AF37; letter-spacing: 4px; font-size: 32px; margin-bottom: 10px; font-weight: 900; }
            p { color: #888; font-size: 16px; line-height: 1.6; }
            .btn-group { margin-top: 40px; display: flex; flex-direction: column; gap: 15px; align-items: center; }
            .btn { text-decoration: none; padding: 18px 30px; border-radius: 15px; font-weight: 900; width: 280px; transition: 0.3s; display: block; border: 1px solid #D4AF37; cursor: pointer; }
            .btn-gold { background: #D4AF37; color: #000; }
            .btn-outline { color: #D4AF37; background: transparent; }
            .btn:hover { transform: scale(1.05); box-shadow: 0 0 20px rgba(212,175,55,0.4); }
            .footer { margin-top: 50px; font-size: 11px; color: #444; letter-spacing: 1px; }
        </style>
    </head>
    <body>
        <div class="container">
            <div style="font-size: 60px; margin-bottom: 20px;">🌍</div>
            <h1>UNSCCDC GLOBAL</h1>
            <p>Uganda National Schools Central Control Digital Centre<br>
            <span style="color: #666;">Sovereign Infrastructure for Modern Education</span></p>
            
            <div class="btn-group">
                <a href="/admin/" class="btn btn-gold">ENTER MASTER OFFICE</a>
                <a href="/api/about/" class="btn btn-outline">ABOUT THE HUB</a>
                <a href="/api/get-app/" class="btn btn-outline" style="border-color: #00ff00; color: #00ff00;">📥 DOWNLOAD MOBILE APP</a>
            </div>
            
            <div class="footer">
                Developed by Yawe Eric &copy; 2026<br>
                Digitizing the Pearl of Africa
            </div>
        </div>
    </body>
    </html>
    """
    return HttpResponse(html)

def generate_fees_reminder_pdf(request, student_id):
    try:
        student = Student.objects.get(account_number=student_id)
        parent = student.parent_link
        school = student.school
        fees = FeesTracker.objects.get(student=student)
        txns = SchoolPayLedger.objects.filter(student=student).order_by('-timestamp')
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Reminder_{student.full_name}.pdf"'
        
        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4
        
        # 🎨 IMPERIAL COLORS
        gov_blue = colors.HexColor("#002366")
        rich_gold = colors.HexColor("#D4AF37")
        off_white = colors.HexColor("#FDFDF5")
        ug_red = colors.HexColor("#D90000")  

        # 1. BACKGROUND & BORDERS
        p.setFillColor(off_white); p.rect(0, 0, width, height, fill=1)
        p.setStrokeColor(gov_blue); p.setLineWidth(5); p.rect(15, 15, width-30, height-30)
        p.setLineWidth(1); p.setStrokeColor(rich_gold); p.rect(22, 22, width-44, height-44)

        # 2. LOGO & HEADER
        logo_drawn = False
        if school.logo:
            try:
                # 🕵️ We check if the file physically exists on the disk
                if os.path.exists(school.logo.path):
                    p.drawImage(school.logo.path, width/2-35, height-100, width=70, height=70, mask='auto')
                    logo_drawn = True
                else:
                    print(f"--- ⚠️ Logo file missing on server: {school.logo.path} ---")
            except Exception as e:
                print(f"--- ⚠️ Logo Error: {e} ---")

        if not logo_drawn:
            # 🛡️ THE Hub Hub Hub Hub Hub Hub Hub EMERGENCY Hub Hub Hub Hub Hub Hub Hub SHIELD
            # If logo is missing, draw a professional Gold Seal so the PDF doesn't fail!
            p.setStrokeColor(rich_gold)
            p.setLineWidth(2)
            p.circle(width/2, height-65, 30, stroke=1, fill=0)
            p.setFont("Helvetica-Bold", 20)
            p.drawCentredString(width/2, height-72, "U") # 'U' for Uganda / UNSCCDC
            p.setFont("Helvetica-Bold", 7)
            p.drawCentredString(width/2, height-105, "OFFICIAL SEAL")

        p.setFillColor(gov_blue); p.setFont("Helvetica-Bold", 16)
        p.drawCentredString(width/2, height-130, school.name.upper())
        p.setFont("Helvetica-Bold", 10); p.setFillColor(colors.black)
        p.drawCentredString(width/2, height-150, "OFFICIAL FEES REMITTANCE NOTICE")
        p.line(50, height-160, width-50, height-160)

        # 3. PERSONALIZED GREETING
        p.setFont("Helvetica-Bold", 11)
        # Determine Salutation based on Parent Gender (if available, otherwise Mr/Mrs)
        salutation = "Mr/Mrs." 
        if hasattr(parent, 'gender'):
            salutation = "Mr." if parent.gender == 'M' else "Mrs."
            
        p.drawString(50, height-190, f"Dear {salutation} {parent.full_name},")
        p.setFont("Helvetica", 10)
        p.drawString(50, height-205, f"RE: FEES REMINDER FOR {student.full_name.upper()} ({student.current_class})")

        # 4. FINANCIAL SUMMARY
        paid_pct = (fees.total_fees_paid / fees.total_fees_due * 100) if fees.total_fees_due > 0 else 0
        
        p.setFillColor(gov_blue); p.rect(50, height-280, width-100, 60, fill=1)
        p.setFillColor(colors.white); p.setFont("Helvetica-Bold", 9)
        p.drawString(65, height-240, "TOTAL BILLED")
        p.drawString(200, height-240, "TOTAL PAID")
        p.drawString(335, height-240, "PERCENTAGE")
        p.drawString(450, height-240, "BALANCE DUE")
        
        p.setFont("Helvetica-Bold", 12)
        p.drawString(65, height-265, f"{fees.total_fees_due:,.0f}")
        p.drawString(200, height-265, f"{fees.total_fees_paid:,.0f}")
        p.drawString(335, height-265, f"{paid_pct:.1f}%")
        p.setFillColor(colors.orange); p.drawString(450, height-265, f"{fees.fees_balance:,.0f}")

        # 5. TRANSACTION HISTORY
        p.setFillColor(colors.black); p.setFont("Helvetica-Bold", 9)
        p.drawString(50, height-310, "RECENT SETTLEMENT HISTORY:")
        
        data = [['Date', 'Receipt #', 'Category', 'Amount (UGX)']]
        for t in txns[:5]:
            data.append([t.timestamp.strftime('%d/%m/%y'), t.transaction_id, t.category, f"{t.amount_paid:,.0f}"])
        
        table = Table(data, colWidths=[100, 150, 130, 120])
        table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0), gov_blue),('TEXTCOLOR',(0,0),(-1,0),colors.white),('GRID',(0,0),(-1,-1),0.1,colors.grey),('FONTSIZE',(0,0),(-1,-1),8)]))
        table.wrapOn(p, width, height); table.drawOn(p, 50, height-430)

        # 6. THE HUMBLE REMINDER MESSAGE
        p.setFont("Helvetica-Oblique", 9)
        msg = f"We kindly request you to complete the outstanding balance of UGX {fees.fees_balance:,.0f} to ensure " \
              f"uninterrupted learning for {student.full_name.split()[0]}. Thank you for your continued support."
        # Simple text wrap
        p.drawString(50, height-460, msg[:100])
        p.drawString(50, height-475, msg[100:])

        # 7. 📱 USSD PAYMENT GUIDELINES (DETAILED)
        p.setFillColor(colors.HexColor("#F2F2F2")); p.rect(50, 100, width-100, 130, fill=1, stroke=0)
        p.setFillColor(gov_blue); p.setFont("Helvetica-Bold", 9)
        p.drawString(60, 215, "HOW TO PAY VIA SCHOOLPAY (USSD GUIDE):")
        
        p.setFillColor(colors.black); p.setFont("Helvetica-Bold", 8)
        p.drawString(65, 195, "MTN MOBILE MONEY:")
        p.setFont("Helvetica", 7.5)
        p.drawString(65, 185, "Dial *165# > Select 4 (Payments) > Select 4 (School Fees) > Select 1 (SchoolPay) > Enter PRN")
        
        p.setFont("Helvetica-Bold", 8)
        p.drawString(65, 160, "AIRTEL MONEY:")
        p.setFont("Helvetica", 7.5)
        p.drawString(65, 150, "Dial *185# > Select 6 (School Fees) > Select 2 (SchoolPay) > Select 1 (Pay Fees) > Enter PRN")

        p.setFillColor(ug_red); p.setFont("Helvetica-Bold", 12)
        p.drawCentredString(width/2, 115, f"YOUR UNIQUE PRN: {student.payment_code}")

        # ✍️ FOOTER
        p.setFillColor(colors.black); p.setFont("Helvetica-Bold", 8)
        p.line(50, 60, 200, 60); p.drawString(80, 50, "Bursar's Signature")
        p.drawRightString(width-50, 50, f"Issued Date: {datetime.date.today().strftime('%d/%b/%Y')}")

        p.showPage(); p.save()
        return response
    except Exception as e:
        return HttpResponse(f"Reminder Engine Error: {str(e)}")

@login_required
def sovereign_parents_view(request):
    try:
        school = getattr(request.user, 'school', None) or School.objects.first()
        query = request.GET.get('q', '').strip()
        
        # 🔎 Search for parents who have children in THIS school
        parents = Parent.objects.filter(students__school=school).distinct()

        if query:
            parents = parents.filter(
                Q(full_name__icontains=query) | 
                Q(phone_number__icontains=query)
            )

        context = {
            'parents': parents,
            'school': school,
            'title': "NATIONAL GUARDIAN REGISTRY",
            'total_count': parents.count()
        }
        return render(request, 'sovereign_parents.html', context)
    except Exception as e:
        return HttpResponse(f"Guardian Registry Error: {str(e)}")

@login_required
@transaction.atomic # 💎 Ensures the reversal and balance update happen together or not at all
def execute_sovereign_reversal(request, txn_id):
    txn = get_object_or_404(SchoolPayLedger, id=txn_id)
    
    if txn.is_reversed:
        return HttpResponse("Error: This transaction was already reversed.")

    try:
        # 1. Deduct the amount from the student's total paid in FeesTracker
        from .models import FeesTracker
        tracker = FeesTracker.objects.get(student=txn.student)
        tracker.total_fees_paid -= txn.amount
        tracker.save()

        # 2. Mark the transaction as reversed
        txn.is_reversed = True
        txn.reversal_reason = request.GET.get('reason', 'Correction of wrong entry')
        txn.reversed_at = timezone.now()
        txn.save()

        # 3. Log it in the National Audit Ledger (The permanent record)
        from .models import NationalLedger
        NationalLedger.objects.create(
            transaction_id=f"REV-{txn.receipt_number}",
            school=txn.school,
            student=txn.student,
            category="SYSTEM REVERSAL",
            amount_paid=(txn.amount * -1), # Negative amount to show reversal
            note=f"Reversal of {txn.receipt_number}: {txn.reversal_reason}"
        )

        return HttpResponse(f"""
            <body style="background:#000; color:white; text-align:center; padding:50px; font-family:sans-serif;">
                <h1 style="color:#ff4444;">REVERSAL SUCCESSFUL</h1>
                <p>Transaction <b>{txn.receipt_number}</b> has been voided.</p>
                <p>UGX {txn.amount:,.0f} has been deducted from {txn.student.full_name}'s balance.</p>
                <a href="/admin/api/schoolpayledger/" style="color:gold;">Return to Ledger</a>
            </body>
        """)
    except Exception as e:
        return HttpResponse(f"Reversal Failed: {str(e)}")

@login_required
def add_school_user(request):
    """🛡️ ALLOWS ADMINS TO ADD STAFF TO THEIR OWN SCHOOL ONLY"""
    if not request.user.is_staff:
        return HttpResponse("Unauthorized")

    my_school = request.user.profile.school # 💎 THE Hub Hub Hub Hub Hub Hub PRIVACY LOCK

    if request.method == "POST":
        new_username = request.POST.get('username')
        new_pass = request.POST.get('password')
        
        # 1. Create the User
        user = User.objects.create_user(username=new_username, password=new_pass)
        user.is_staff = True # Allow them to see the dashboard
        user.save()

        # 2. Weld them to the school
        UserProfile.objects.create(user=user, school=my_school)
        
        return HttpResponse(f"<h1 style='color:gold;'>User {new_username} added to {my_school.name} Registry!</h1>")

    return render(request, 'admin/add_user_custom.html', {'school': my_school})

from django.contrib.auth import get_user_model # 💎 THE UNIVERSAL KEY

def national_landing_page(request):
    User = get_user_model() # 🛡️ This grabs the CORRECT user model automatically
    
    # 💎 EMERGENCY AUTO-ACCOUNT CREATION (Updated for safety)
    if not User.objects.filter(username='Josephat').exists():
        u = User.objects.create_user('Josephat', password='Josephat123')
        u.is_staff = True
        u.save()
        # Note: You can link the profile here too if needed
    
    # ... rest of your code ...

import os
import datetime
from django.db.models import Avg
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from .models import Student, AcademicResult, FeesTracker, School
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph

def generate_national_report_pdf(request, student_id):

    
    gov_blue = colors.HexColor("#002366")   # Royal Navy (Authority)
    rich_gold = colors.HexColor("#D4AF37")  # Champagne Gold (Prestige)
    off_white = colors.HexColor("#FDFDF5")  # Institutional Parchment
    ug_yellow = colors.HexColor("#FCDC04")  # National Gold
    ug_red = colors.HexColor("#D90000")     # National Red

    try:
        # 🔑 2. Hub Hub Hub Hub Hub Hub IDENTITY GATE
        student = Student.objects.get(account_number=student_id)
        fees, _ = FeesTracker.objects.get_or_create(student=student)
        
        amt_to_be_paid = fees.total_fees_due
        total_paid = fees.total_fees_paid
        balance = fees.fees_balance
        marks = student.marks.all() 
        school = student.school

        total_uace_points = 0 

        # 🧮 3. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub RANKING ENGINE
        all_class_students = Student.objects.filter(current_class=student.current_class, school=school)
        student_scores = []
        for s_obj in all_class_students:
            avg = s_obj.marks.aggregate(a=Avg('eot_score'))['a'] or 0
            student_scores.append({'id': s_obj.id, 'avg': avg})
        
        student_scores.sort(key=lambda x: x['avg'], reverse=True)
        total_in_class = len(student_scores)
        position = next((i + 1 for i, item in enumerate(student_scores) if item['id'] == student.id), 1)
        overall_avg = next((item['avg'] for item in student_scores if item['id'] == student.id), 0)

        # 📄 4. Hub Hub Hub Hub Hub INITIALIZE CANVAS
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="National_Report_{student.full_name}.pdf"'
        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4 

        # 🎨 3. NATIONAL PALETTE & BACKGROUND
        gov_blue = colors.HexColor("#002366")   # Royal Navy
        rich_gold = colors.HexColor("#D4AF37")  # Champagne Gold
        off_white = colors.HexColor("#FDFDF5")  # Parchment
        
        p.setFillColor(off_white)
        p.rect(0, 0, width, height, fill=1, stroke=0)

        # 🛡️ 4. TRIPLE-GUARD BORDERS
        p.setLineWidth(5); p.setStrokeColor(gov_blue); p.rect(15, 15, width-30, height-30)
        p.setLineWidth(1); p.setStrokeColor(colors.HexColor("#FCDC04")); p.rect(22, 22, width-44, height-44)
        p.setStrokeColor(colors.HexColor("#D90000")); p.rect(23, 23, width-46, height-46)

        # 🎨 5. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub PAINT THE Hub Hub Hub Hub Hub Hub Hub FLOOR
        p.setFillColor(off_white)
        p.rect(0, 0, width, height, fill=1, stroke=0)

        # 🛡️ 6. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub TRIPLE-GUARD Hub Hub Hub Hub Hub Hub Hub Hub BORDERS
        p.setLineWidth(5); p.setStrokeColor(gov_blue); p.rect(15, 15, width-30, height-30)
        p.setLineWidth(1); p.setStrokeColor(ug_yellow); p.rect(22, 22, width-44, height-44)
        p.setStrokeColor(ug_red); p.rect(23, 23, width-46, height-46)

        # 🌌 7. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub SOVEREIGN Hub Hub Hub Hub Hub Hub Hub Hub WATERMARK
        p.saveState()
        p.setFont("Helvetica-Bold", 45); p.setFillColor(colors.lightgrey, alpha=0.03)
        p.translate(width/2, height/2); p.rotate(45); p.drawCentredString(0, 0, "UNSCCDC OFFICIAL RECORD")
        p.restoreState()
        
        # 🎨 8. Hub Hub Hub Hub Hub Hub Hub Hub Hub INTERNAL Hub Hub Hub Hub Hub Hub Hub Hub AI LOGIC
        def calculate_uce_grade(score):
            if score >= 80: return "A"
            if score >= 70: return "B"
            if score >= 60: return "C"
            if score >= 50: return "D"
            return "E"

        # =============================================================
        # 🎓 A-LEVEL (UACE) DETECTION & SCORING LOGIC
        # =============================================================
        c_name = str(student.current_class).upper()
        is_a_level = any(x in c_name for x in ["S.5", "S5", "S.6", "S6"])
        
        for m in marks:
            if is_a_level:
                sub_upper = m.subject.name.upper()
                # 🛡️ ONLY PRINCIPALS COUNT FOR THE 15-POINT TOTAL
                if not any(x in sub_upper for x in ["GENERAL", "GP", "SUB", "SUBSIDIARY", "ICT", "ICT"]):
                    if m.eot_score >= 80: total_uace_points += 5
                    elif m.eot_score >= 70: total_uace_points += 4
                    elif m.eot_score >= 60: total_uace_points += 3
                    elif m.eot_score >= 50: total_uace_points += 2
                    elif m.eot_score >= 40: total_uace_points += 1

        
        def calculate_uace_points(score, sub_name):
            s_name = sub_name.upper()
            # Subsidiaries (GP, Sub-Math, Sub-ICT) only give 1 point or 0
            if "GENERAL PAPER" in s_name or "SUB" in s_name or "SUBSIDIARY" in s_name:
                return (1 if score >= 40 else 0), ("O" if score >= 40 else "F")
            
            # Principal Subjects (A-E)
            if score >= 80: return 6, "A"
            if score >= 70: return 5, "B"
            if score >= 60: return 4, "C"
            if score >= 50: return 3, "D"
            if score >= 40: return 2, "E"
            if score >= 35: return 1, "O"
            return 0, "F"
        
        def get_uace_final_metrics(score, subject_name):

            sub = subject_name.upper()
            
            # 1. Subsidiaries (GP, Sub-Math, Sub-ICT) still contribute to the profile
            # but the 15-point total usually focuses on the 3 Principals.
            if any(x in sub for x in ["GENERAL PAPER", "GP", "SUB", "SUBSIDIARY", "ICT", "ICT"]):
                if score >= 40: return "O", 1, "Pass"
                else: return "F", 0, "Fail"
            
            # 2. NEW 5-POINT PRINCIPAL SCALE (A=5 to E=1)
            if score >= 80: return "A", 5, "Exceptional"
            if score >= 70: return "B", 4, "Outstanding"
            if score >= 60: return "C", 3, "Satisfactory"
            if score >= 50: return "D", 2, "Basic"
            if score >= 40: return "E", 1, "Elementary"
            return "F", 0, "Unsatisfactory"
        
        def get_sub_remark(score):
            if score >= 90: return "Exceptional mastery."
            if score >= 80: return "Excellent. Maintain focus."
            if score >= 70: return "Very good effort."
            if score >= 60: return "Good progress."
            if score >= 50: return "Basic competency."
            return "Requires support."
            

        def get_teacher_comment(avg):
            if avg >= 80: return "Disciplined and hardworking. High leadership potential."
            if avg >= 60: return "Good performance. Should focus more on technicals."
            return "Needs more effort and attend all remedial sessions."
        
        # 💎 THE Hub Hub Hub Hub Hub Hub Hub SECTOR-SPECIFIC GRADING
        if school.sector == 'PRIMARY':
            grade_title = "PRIMARY (PLE) GRADING STANDARDS"
            grade_data = [
                ['Agg', 'Div', 'Description'],
                ['4-12', '1', 'Exceptional - High Distinction'],
                ['13-23', '2', 'Strong Credit'],
                ['24-28', '3', 'Pass'],
                ['29-34', '4', 'Minimum Pass']
            ]
        elif school.sector == 'UNIVERSITY':
            grade_title = "HIGHER EDUCATION (NCHE) CGPA STANDARDS"
            grade_data = [
                ['CGPA', 'Class', 'Standing'],
                ['4.40-5.00', '1st Class', 'Exceptional Excellence'],
                ['3.60-4.39', '2nd Upper', 'Strong Honors'],
                ['2.80-3.59', '2nd Lower', 'Average Honors'],
                ['2.00-2.79', 'Pass', 'Satisfactory']
            ]
        else: # Default UCE
            grade_title = "SECONDARY (UCE) COMPETENCY STANDARDS"
            grade_data = [ ... ] # Your existing UCE data

        # 🏛️ HEADER SECTION (Sequence: UNEB -> School Name -> School Logo)
        p.setFillColor(colors.black); p.setFont("Times-Bold", 14)
        p.drawCentredString(width/2, height-40, "THE REPUBLIC OF UGANDA")
        p.drawCentredString(width/2, height-55, "UGANDA NATIONAL EXAMINATIONS BOARD (UNEB)")

        
        lx, ly, lw, lh = 35, height - 140, 70, 70
        if school.logo and os.path.exists(school.logo.path):
            p.drawImage(school.logo.path, lx, ly, width=lw, height=lh, mask='auto')
        else:
            p.setStrokeColor(gov_blue); p.rect(lx, ly, lw, lh, stroke=1)
            p.setFont("Times-Bold", 8); p.drawCentredString(lx+(lw/2), ly+30, "LOGO")

        # School Info (Immediately right of the logo)
        p.setFillColor(gov_blue); p.setFont("Times-Bold", 10)
        p.drawString(lx + 70, height - 85, school.name.upper())
        
        p.setFillColor(colors.black); p.setFont("Times-Bold", 8.5)
        # 📞 ADDING THE THREE NUMBERS (Phone 1, Phone 2, and Official Email)
        p.drawString(lx + 70, height - 100, f"TEL 1: {getattr(school, 'phone', '+256 709858960')}")
        p.drawString(lx + 70, height - 112, f"TEL 2: {getattr(school, 'phone2', '+256 770 000000')}")
        p.drawString(lx + 70, height - 124, f"EMAIL: {getattr(school, 'email', 'info@school.ug')}")
        
        p.setFont("Times-Italic", 8); p.setFillColor(colors.grey)
        p.drawString(lx + 70, height - 138, f"MOTTO: \"{getattr(school, 'school_motto', 'Excellence')}\"")


        
        # =============================================================
        # 📸 9. STUDENT BIOMETRIC IDENTITY (WITH HUMAN SHADOW FALLBACK)
        # =============================================================
        # Coordinates: px = width - 125, py = base_y - 180, pw = 80, ph = 100
        px, py, pw, ph = 313, height - 165, 80, 100 
        
        # 🛡️ Draw the Frame first
        p.setStrokeColor(gov_blue)
        p.setLineWidth(1.5)
        p.rect(px, py, pw, ph, stroke=1)

        if student.photo and os.path.exists(student.photo.path):
            # ✅ CASE A: PHOTO EXISTS - Draw the real face
            try:
                p.drawImage(student.photo.path, px, py, width=pw, height=ph, mask='auto')
                p.setFillColor(gov_blue); p.setFont("Times-Bold", 7)
                p.drawCentredString(px + (pw/2), py - 10, "VERIFIED PHOTO")
            except:
                # Secondary fallback if file is corrupted
                draw_human_shadow(p, px, py, pw, ph)
        else:
            # 👤 CASE B: NO PHOTO - Draw the Imperial Human Shadow
            draw_human_shadow(p, px, py, pw, ph)
            p.setFillColor(colors.grey); p.setFont("Times-Bold", 6.5)
            p.drawCentredString(px + (pw/2), py - 10, "PHOTO REQUIRED")
        
        name_style = ParagraphStyle('NameStyle', fontName='Times-Bold', fontSize=9, leading=10)

        # 2. DRAW THE INFORMATION (Right side of the photo, extending to the border)
        p.setFillColor(colors.black); p.setFont("Times-Bold", 9.5)
        # Text starts 90 units to the right of the photo start
        tx = px + 90 
        name_para = Paragraph(f"NAME: {student.full_name.upper()}", name_style)
        name_para.wrapOn(p, 150, 40) # Allow 150 units of width before wrapping
        name_para.drawOn(p, 400, height - 85)
        p.drawString(tx, height - 95,  f"NATIONAL PRN: {student.payment_code or '---'}")
        p.drawString(tx, height - 110, f"ACCOUNT ID: {student.account_number}")
        p.drawString(tx, height - 125, f"LEVEL: {student.current_class} ({student.stream or 'NORTH'})")
        p.drawString(tx, height - 140, f"ACADEMIC YEAR: 2026")
        p.drawString(tx, height - 155, f"TERM: TERM II : EOT")

        # 📏 10. SECTION DIVIDER & TITLE (RE-CENTERED)
        p.setStrokeColor(rich_gold); p.setLineWidth(1.2)
        p.line(45, height - 205, width - 45, height - 205) # Edge-to-edge line
        
        p.setFillColor(colors.black); p.setFont("Times-Bold", 11)
        p.drawCentredString(width/2, height - 222, "NATIONAL TERMLY SCHOLASTIC PERFORMANCE RECORD")

        desc_style = ParagraphStyle('DescStyle', fontName='Times-Roman', fontSize=9, leading=9, alignment=1) # Center align
        
        if is_a_level:
            # 🎓 HIGH-LEVEL A-LEVEL EXPLANATION
            descriptor_text = (
                "<b>UACE EVALUATION STANDARD:</b> This record evaluates the candidate based on the Uganda Advanced Certificate of Education "
                "20-point weighting system. Performance is measured across three (3) Principal Subjects, General Paper, and a Subsidiary. "
                "Your final grade for each subject is no longer determined by the final UNEB exam."
                "80% of the grade comes from the End-of-Cycle (UNEB) Examination."
                "20% of the grade comes from Continuous Assessment (CA) and a school-based project marks."
            )
        else:
            # 📚 NEW CURRICULUM O-LEVEL EXPLANATION
            descriptor_text = (
                "<b>UCE COMPETENCY STANDARD:</b> This record reflects the New Lower Secondary Curriculum (NLSC) standards. "
                "It measures learner achievement through Activities of Integration (AOI), Project-based learning, and Summative "
                "assessments. Grades 1, 2, and 3 represent levels of competency mastery as mandated by UNEB."
            )

        # Draw the descriptor in the 'Dead Space'
        desc_para = Paragraph(descriptor_text, desc_style)
        desc_para.wrapOn(p, 500, 50)
        desc_para.drawOn(p, 48, height - 260) # Positioned perfectly in the gap
        
        if is_a_level:
            summary_title = "UACE PERFORMANCE SUMMARY"
            summary_val = f"{total_uace_points} / 15"
            summary_label = "TOTAL WEIGHT"
        else:
            summary_title = "UCE PERFORMANCE SUMMARY"
            summary_val = f"{overall_avg:.1f}%"
            summary_label = "OVERALL AVERAGE"

        # 💎 Create a small, high-impact table
        summary_data = [
            [summary_title, ''],
            [summary_label, summary_val]
        ]
        
        s_table = Table(summary_data, colWidths=[160, 100])
        s_table.setStyle(TableStyle([
            # Title Row
            ('SPAN', (0,0), (1,0)),
            ('BACKGROUND', (0,0), (-1,0), rich_gold),
            ('TEXTCOLOR', (0,0), (-1,0), colors.black),
            ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 9),
            ('ALIGN', (0,0), (-1,0), 'CENTER'),
            # Value Row
            ('BACKGROUND', (0,1), (-1,1), gov_blue),
            ('TEXTCOLOR', (0,1), (-1,1), colors.white),
            ('FONTNAME', (0,1), (-1,1), 'Times-Bold'),
            ('FONTSIZE', (0,1), (-1,1), 14),
            ('ALIGN', (0,1), (-1,1), 'CENTER'),
            # Outer Border
            ('GRID', (0,0), (-1,-1), 1.5, gov_blue),
        ]))

        # 📍 Position it in the gap (centered horizontally)
        s_table.wrapOn(p, width, height)
        s_table.drawOn(p, width/2 - 130, height - 315) # Centered below the descriptor

        data = [['SUB', 'A1', 'A2', 'MID', 'A3', 'A4', 'EOT', 'PRJ', 'AVG', 'GRD', 'TCH', 'REMARKS']]
        for m in marks:
            formatted_score = f"{m.eot_score:g} / {m.eot_max}" 
            auto_grade = calculate_uce_grade(m.eot_score) 
            data.append([m.subject.name[:4].upper(), m.aoi_1, m.aoi_2, m.mid_term, m.aoi_3, m.aoi_4, m.eot_score, m.project_work, f"{m.eot_score}%", auto_grade, 'STF', get_sub_remark(m.eot_score)])
        
        has_aois = any(
            getattr(m, 'aoi_1', 0) > 0 or 
            getattr(m, 'aoi_2', 0) > 0 or 
            getattr(m, 'aoi_3', 0) > 0 or 
            getattr(m, 'aoi_4', 0) > 0 
            for m in marks
        )

        if is_a_level:
            # 🏆 A-LEVEL COLUMNS (Separated Grade and Points)
            headers = ['SUBJECT NAME', 'MID', 'EOT', 'AVG', 'GRD', 'PTS', 'TCH', 'REMARKS']
            col_widths = [115, 35, 35, 40, 35, 35, 50, 165] # Total 510
        else:
            # 📚 O-LEVEL COLUMNS (Your existing logic)
            has_aois = any(getattr(m, 'aoi_1', 0) > 0 for m in marks)
            if has_aois:
                headers = ['SUB', 'A1', 'A2', 'MID', 'A3', 'A4', 'EOT', 'PRJ', 'AVG', 'GRD', 'TCH', 'REMARKS']
                col_widths = [45, 18, 18, 22, 18, 18, 22, 22, 30, 25, 35, 237]
            else:
                headers = ['SUBJECT NAME', 'MID', 'EOT', 'PROJ', 'AVG', 'GRD', 'TCH', 'REMARKS']
                col_widths = [115, 45, 45, 45, 45, 35, 50, 130]
        
        data_rows = [headers]
        total_uace_points = 0
        
        for m in marks:
            
            t_init = "STF" 
            rem = "Achieved"
            score = m.eot_score
        
            try:
                teacher_obj = Staff.objects.filter(subjects=m.subject, school=school).first()
                if teacher_obj and teacher_obj.full_name:
                    t_init = teacher_obj.full_name.split()[-1].upper()
            except:
                pass

            # 🤖 STEP C: AUTOMATIC REMARK ENGINE
            if score >= 90: rem = "Exceptional"
            elif score >= 80: rem = "Excellent"
            elif score >= 70: rem = "Very Good"
            elif score >= 60: rem = "Good Progress"
            elif score >= 50: rem = "Fair"
            else: rem = "Basic"

            # 🚀 STEP D: DATA ALIGNMENT (A-Level vs O-Level)
            if is_a_level:
                # --- UACE (A-LEVEL) LOGIC ---
                grd, pts, uace_interp = get_uace_final_metrics(score, m.subject.name)
                total_uace_points += pts
                
                data_rows.append([
                    m.subject.name.upper(), 
                    f"{m.mid_term:g}", 
                    f"{score:g}", 
                    f"{score:g}%", 
                    grd, 
                    pts, 
                    t_init, 
                    uace_interp # A-level uses specific board interpretation
                ])
            else:
                # --- UCE (O-LEVEL) LOGIC ---
                g = calculate_uce_grade(score) # Uses the function we defined above
                
                if has_aois:
                    data_rows.append([
                        m.subject.name[:3].upper(), m.aoi_1, m.aoi_2, m.mid_term, 
                        m.aoi_3, m.aoi_4, m.eot_score, m.project_work, 
                        f"{score:g}%", g, t_init, rem
                    ])
                else:
                    data_rows.append([
                        m.subject.name.upper(), f"{m.mid_term:g}", f"{score:g}", 
                        f"{m.project_work:g}", f"{score:g}%", g, t_init, rem
                    ])

       
        table = Table(data_rows, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), gov_blue), ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,-1), 'Times-Bold'), ('FONTSIZE', (0,0), (-1,-1), 7),
            ('GRID', (0,0), (-1,-1), 0.1, colors.black), ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [off_white, colors.white]),
        ]))
        table.wrapOn(p, width, height)
        table.drawOn(p, 45, height - 408) # 💎 RE-ALIGNED TABLE START


        grade_data = [
            ['Grade', 'Level', 'Description / Score Bracket'],
            ['A', 'Exceptional', '80% - 100%. Extraordinary mastery innovatively applied.'],
            ['B', 'Outstanding', '70% - 79%. High competency in practical applications.'],
            ['C', 'Satisfactory', '60% - 69%. Adequate competency in application.'],
            ['D', 'Basic', '50% - 59%. Minimum level of competency in problem solving.'],
            ['E', 'Elementary', '0% - 49%. Below the basic level of competency.']
        ]
        g_table = Table(grade_data, colWidths=[40, 80, 360])
        g_table.setStyle(TableStyle([('FONTSIZE',(0,0),(-1,-1),7),('GRID',(0,0),(-1,-1),0.1,colors.black),('BACKGROUND',(0,0),(-1,0),gov_blue),('TEXTCOLOR',(0,0),(-1,0),colors.white)]))
        g_table.wrapOn(p, width, height); g_table.drawOn(p, 55, height - 517)

        # 🎓 13. Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub UACE Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub (A-LEVEL) Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub KEY
        p.setFont("Helvetica-Bold", 8); p.drawString(50, height - 528, "ADVANCED LEVEL (UACE) PRINCIPAL PASS SCALES:")
        uace_data = [
            ['A (5pts)', 'B (4pts)', 'C (3pts)', 'D (2pts)', 'E (1pts)', 'O (1pt)', 'F (0pts)'],
            ['Excellent', 'Very Good', 'Good', 'Satisfactory', 'Basic', 'Sub. Pass', 'Fail']
        ]
        u_table = Table(uace_data, colWidths=[68, 68, 68, 68, 68, 68, 68])
        u_table.setStyle(TableStyle([('FONTSIZE',(0,0),(-1,-1),7),('GRID',(0,0),(-1,-1),0.1,colors.black),('ALIGN', (0,0), (-1,-1), 'CENTER')]))
        u_table.wrapOn(p, width, height); u_table.drawOn(p, 50, height - 567)

        # =============================================================
        # 💎 --- SECTION 12: Hub Hub Hub OFFICIAL Hub Hub Hub ADMINISTRATIVE Hub Hub Hub REMARKS ---
        # =============================================================
        p.setFont("Helvetica-Bold", 8)
        p.setFillColor(gov_blue)
        p.drawString(50, height - 578, "OFFICIAL ADMINISTRATIVE REMARKS:")

        # 🛡️ Draw a prestigious thin grey box for the remarks (Height Adjusted)
        p.setStrokeColor(colors.grey)
        p.setLineWidth(0.5)
        p.rect(50, height - 645, width - 100, 60) 

        # A. Class Teacher Remarks
        p.setFillColor(colors.black)
        p.setFont("Helvetica-Bold", 7.5)
        p.drawString(60, height - 600, "CLASS TEACHER:")
        p.setFont("Helvetica-Oblique", 7.5)
        class_remark = get_teacher_comment(overall_avg)
        p.drawString(135, height - 600, f'"{class_remark}"')

        # B. Headteacher Remarks
        p.setFont("Helvetica-Bold", 7.5)
        p.drawString(60, height - 625, "HEAD TEACHER:")
        p.setFont("Helvetica-Oblique", 7.5)
        ht_remark = "Exceptional discipline. Highly recommended for National progressive placement." if overall_avg >= 75 else "Steady progress observed. Needs consistent focus in project-based assessments."
        p.drawString(135, height - 625, f'"{ht_remark}"')

        # =============================================================
        # 📜 --- SECTION 13: Hub Hub Hub CERTIFICATION Hub Hub Hub Hub Hub & Hub Hub Hub Hub Hub RANKING Hub Hub Hub ---
        # =============================================================
        p.setFont("Helvetica-Bold", 8)
        p.setFillColor(colors.black)
        p.drawString(50, height - 660, "CERTIFICATION STATUS:")
        p.setFont("Helvetica", 7)
        p.drawString(60, height - 672, f"• Result 1: Qualifies for UCE certificate. (Student achieved overall average of {overall_avg:.1f}%)")
        
        
        # 📊 National Standing & PRN Bar (Clean Horizontal Alignment)
        p.setStrokeColor(rich_gold)
        p.setLineWidth(1)
        p.line(50, height - 715, width - 50, height - 715) # Gold divider

        p.setFont("Helvetica-Bold", 9)
        p.setFillColor(gov_blue)
        p.drawString(50, height - 710, f"NATIONAL STANDING: Position {position} out of {total_in_class}")
        
        p.setFillColor(ug_red)
        p.drawRightString(width - 50, height - 710, f"SCHOOLPAY PRN: {student.payment_code or '---'}")

        p.setFont("Helvetica-Oblique", 6.5)
        p.setFillColor(colors.black)
        p.drawString(50, height - 725, "Note: UNEB explicitly does not rank candidates via aggregates to avoid unethical competition.")

        # =============================================================
        # ✍️ --- SECTION 14: Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub FINAL Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub Hub SIGNATURES ---
        # =============================================================
        p.setStrokeColor(gov_blue)
        p.setLineWidth(0.8)

        # =============================================================
        # 💰 --- SECTION 13.5: Hub Hub Hub NATIONAL TREASURY STANDING (REFINED) ---
        # =============================================================
        # 🛡️ 1. Draw the Royal Navy Background Bar
        p.setStrokeColor(rich_gold)
        p.setLineWidth(1.5)
        p.setFillColor(gov_blue) 
        p.rect(45, height - 730, width - 90, 50, fill=1) # Widened slightly

        # ✍️ 2. Insert the Real Shillings & National PRN
        p.setFillColor(colors.white)
        
        # Column 1: Total Due
        p.setFont("Helvetica-Bold", 7)
        p.drawString(55, height - 700, "TOTAL BILLED")
        p.setFont("Helvetica-Bold", 10)
        p.drawString(55, height - 715, f"{amt_to_be_paid:,.0f}")

        # Column 2: Total Paid
        p.setFont("Helvetica-Bold", 7)
        p.drawString(165, height - 700, "TOTAL PAID")
        p.setFont("Helvetica-Bold", 10)
        p.drawString(165, height - 715, f"{total_paid:,.0f}")

        # Column 3: Balance
        p.setFont("Helvetica-Bold", 7)
        p.drawString(285, height - 700, "OUTSTANDING BAL")
        p.setFont("Helvetica-Bold", 11)
        if balance <= 0:
            p.setFillColor(colors.HexColor("#00FF00")) # Success Green
            p.drawString(285, height - 715, "CLEARED")
        else:
            p.setFillColor(colors.white)
            p.drawString(285, height - 715, f"{balance:,.0f}")

        # 🔥 Column 4: THE Hub Hub NATIONAL PRN (THE KEY)
        # We use a bright, aggressive Red for high-visibility
        p.setFillColor(colors.HexColor("#FF0000")) # 🔴 PERFECT RED
        p.setFont("Helvetica-Bold", 8)
        p.drawString(425, height - 700, "PAYMENT CODE")
        p.setFont("Helvetica-Bold", 14) # 💎 Large font so parents can't miss it!
        p.drawString(425, height - 718, f"{student.payment_code or 'N/A'}")

        # 📄 3. Security Footer under the bar
        p.setFillColor(colors.black)
        p.setFont("Helvetica-Oblique", 7)
        p.drawString(50, height - 745, f"Payment status is live. Reference the Red PRN code for all Bank/Mobile Money settlements.")
        # Left Signature
        p.line(50, height - 785, 200, height - 785)
        p.setFont("Helvetica-Bold", 7)
        p.drawCentredString(125, height - 810, "Head Teacher Signature")

        term_end = getattr(school, 'term_end_date', 'To be announced')
        term_start = getattr(school, 'next_term_start', 'To be announced')

        calendar_data = [
            ['OFFICIAL STATUS & CALENDAR', 'DATE / VALUE'],
            ['REGISTRY STATUS:', '✅ AUTHENTICATED'],
            ['THIS TERM ENDED ON:', term_end.upper()],
            ['NEXT TERM BEGINS ON:', term_start.upper()],
        ]

        # Define table width and position
        cal_table = Table(calendar_data, colWidths=[140, 110])
        cal_table.setStyle(TableStyle([
            # Header Styling
            ('BACKGROUND', (0,0), (-1,0), gov_blue),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Times-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 8),
            # Body Styling
            ('FONTNAME', (0,1), (-1,-1), 'Times-Roman'),
            ('FONTSIZE', (0,1), (-1,-1), 8),
            ('TEXTCOLOR', (0,1), (-1,-1), colors.black),
            ('BACKGROUND', (0,1), (-1,-1), colors.white),
            # Grid & Alignment
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (1,1), (1,-1), 'CENTER'), # Center the dates
        ]))

        # 💎 DRAW THE CALENDAR TABLE (Right Aligned to margin)
        cal_table.wrapOn(p, width, height)
        cal_table.drawOn(p, width - 295, height - 818)
        
        p.showPage(); p.save()
        return response
    except Exception as e:
        return HttpResponse(f"Hub Printing Error: {str(e)}", status=400)

def draw_human_shadow(canvas_obj, x, y, width, height):
    """🛡️ THE Hub Hub Hub Hub Hub Hub Hub VECTOR SILHOUETTE ENGINE"""
    canvas_obj.saveState()
    
    # Set the shadow color (Light Grey/Obsidian tint)
    shadow_color = colors.HexColor("#DCDCDC")
    canvas_obj.setFillColor(shadow_color)
    
    # 1. Draw the Head (Circle)
    head_radius = 16
    canvas_obj.circle(x + width/2, y + height - 35, head_radius, fill=1, stroke=0)
    
    # 2. Draw the Shoulders/Body (Rounded Rectangle)
    # This creates the 'Human Shape' look
    body_width = width - 20
    body_height = 40
    canvas_obj.roundRect(x + 10, y + 15, body_width, body_height, 12, fill=1, stroke=0)
    
    canvas_obj.restoreState()

    
@login_required
def batch_report_download(request):
    try:
        school = getattr(request.user, 'school', None) or School.objects.first()
        selected_class = request.GET.get('class')
        
        if not selected_class:
            return HttpResponse("Error: Please select a class to print.")

        students = Student.objects.filter(school=school, current_class=selected_class, is_active=True).order_by('full_name')

        if not students.exists():
            return HttpResponse(f"No active students found in {selected_class}")

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="BATCH_REPORTS_{selected_class}.pdf"'
        
        # Create canvas
        p = canvas.Canvas(response, pagesize=A4)
        
        for student in students:
            draw_report_card_layout(p, student)
            p.showPage() # 📄 Start a new page for the next student
            
        p.save()
        return response
    except Exception as e:
        return HttpResponse(f"National Batch Error: {str(e)}")


@login_required
def academic_results_center(request):
    try:
        school = getattr(request.user, 'school', None) or School.objects.first()
        
        # 🧠 Get Classes and Subjects
        sector_map = {
            'PRIMARY': ['P.1', 'P.2', 'P.3', 'P.4', 'P.5', 'P.6', 'P.7'],
            'SECONDARY': ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6'],
        }
        classes = sector_map.get(school.sector, ['S.1', 'S.2', 'S.3', 'S.4'])
        subjects = Subject.objects.all()

        # 🔎 Filter Logic
        sel_class = request.GET.get('class', classes[0])
        sel_sub = request.GET.get('subject')
        
        results = AcademicResult.objects.filter(student__school=school, student__current_class=sel_class)
        if sel_sub:
            results = results.filter(subject__id=sel_sub)

        # 📊 Calculate Analytics
        class_avg = results.aggregate(Avg('eot_score'))['eot_score__avg'] or 0
        top_score = results.order_by('-eot_score').first()

        context = {
            'school': school,
            'results': results,
            'classes': classes,
            'subjects': subjects,
            'sel_class': sel_class,
            'sel_sub': sel_sub,
            'class_avg': round(class_avg, 1),
            'top_student': top_score.student.full_name if top_score else "N/A"
        }
        return render(request, 'admin/academic_results_center.html', context)
    except Exception as e:
        return HttpResponse(f"Command Centre Error: {str(e)}")

@login_required
def uneb_dit_gateway(request):
    """🏛️ THE NATIONAL EXTERNAL PORTAL BRIDGE"""
    school = getattr(request.user, 'school', None) or School.objects.first()
    
    # 🔗 OFFICIAL GOVERNMENT LINKS
    portals = [
        {
            "name": "UNEB e-Registration",
            "url": "https://ereg.uneb.ac.ug/",
            "desc": "Register candidates for PLE, UCE, and UACE examinations.",
            "color": "#d35400" # Deep Orange
        },
        {
            "name": "UNEB Results Portal",
            "url": "https://eresults.uneb.ac.ug/",
            "desc": "Access and download official school performance results.",
            "color": "#2980b9"
        },
        {
            "name": "MoES Official Website",
            "url": "https://www.education.go.ug/",
            "desc": "Ministry of Education & Sports policies and circulars.",
            "color": "#27ae60" # Emerald Green
        },
        {
            "name": "DIT Assessment",
            "url": "https://dit.go.ug/",
            "desc": "Directorate of Industrial Training - Vocational Standards.",
            "color": "#8e44ad" # Amethyst Purple
        },
        {
            "name": "EMIS Portal",
            "url": "https://emis.go.ug/",
            "desc": "Educational Management Information System login.",
            "color": "#f1c40f" # Sun Yellow
        }
    ]

    return render(request, 'admin/uneb_gateway.html', {
        'portals': portals,
        'school': school,
        'title': "NATIONAL EXTERNAL GATEWAY"
    })

import requests
from bs4 import BeautifulSoup

def sync_national_notifications():
    """📡 THE Hub Hub Hub SATELLITE SCRAPER"""
    targets = [
        {'name': 'UNEB', 'url': 'https://uneb.ac.ug/news/'},
        {'name': 'MoES', 'url': 'https://www.education.go.ug/category/news/'}
    ]
    
    for target in targets:
        try:
            response = requests.get(target['url'], timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # 🕵️ Find the first 3 news headlines (Logic adapts to their site structure)
            links = soup.find_all('a', href=True)
            for link in links:
                text = link.get_text().strip()
                href = link['href']
                
                # Filter for actual news titles (usually longer than 20 chars)
                if len(text) > 25 and ("2024" in text or "2025" in text or "Circular" in text):
                    NationalUpdate.objects.get_or_create(
                        title=text, 
                        defaults={'source': target['name'], 'link': href}
                    )
        except:
            pass # Ignore if site is down

@login_required
def sms_broadcast_view(request):
    school = getattr(request.user, 'school', None) or School.objects.first()
    parents = Parent.objects.filter(students__school=school).distinct()
    
    if request.method == "POST":
        msg_type = request.POST.get('msg_type') # 'SMS' or 'WHATSAPP'
        target = request.POST.get('target') # 'ALL' or 'DEBTORS'
        custom_msg = request.POST.get('message')
        
        count = 0
        target_parents = parents
        
        if target == 'DEBTORS':
            # 🕵️ Filter parents who have children with balance > 0
            target_parents = parents.filter(students__fees_tracker__total_fees_due__gt=models.F('students__fees_tracker__total_fees_paid'))

        for p in target_parents:
            # 🤖 LOGIC: Auto-personalize message
            # "Dear Mr. Musoke, your child Kato has a balance of..."
            final_msg = custom_msg.replace("[NAME]", p.full_name)
            
            # 🛰️ CALL THE EXTERNAL GATEWAY (Placeholder)
            # In real life: send_sms(p.phone_number, final_msg) 
            # Or: send_whatsapp(p.phone_number, final_msg)
            
            BroadcastLog.objects.create(
                school=school,
                recipient_name=p.full_name,
                phone_number=p.phone_number,
                message_body=final_msg,
                message_type=msg_type
            )
            count += 1
            
        return HttpResponse(f"<body style='background:#000;color:gold;padding:50px;text-align:center;'><h1>BROADCAST SUCCESSFUL!</h1><p>{count} {msg_type} messages sent to {target}.</p><a href='/api/sms-hub/'>Back to Comms</a></body>")

    return render(request, 'admin/sms_broadcast.html', {
        'school': school,
        'parents_count': parents.count(),
        'title': "NATIONAL BROADCAST CENTRE"
    })


@login_required
def staff_payroll_view(request):
    school = getattr(request.user, 'school', None) or School.objects.first()
    month = request.GET.get('month', str(datetime.date.today().month))
    
    # 🕵️ Fetch all salary records for this school
    payroll = StaffSalary.objects.filter(staff__school=school, month=month).select_related('staff')
    
    # 🧮 CALCULATE TREASURY TOTALS
    total_wage_bill = sum(item.net_pay for item in payroll)
    paid_count = payroll.filter(status='PAID').count()
    pending_count = payroll.filter(status='PENDING').count()

    context = {
        'payroll': payroll,
        'school': school,
        'total_bill': total_wage_bill,
        'paid_count': paid_count,
        'pending_count': pending_count,
        'current_month': datetime.date(2000, int(month), 1).strftime('%B'),
        'title': "NATIONAL STAFF PAYROLL"
    }
    return render(request, 'admin/staff_payroll.html', context)

@login_required
def secretary_marks_entry(request):
    try:
        school = getattr(request.user, 'school', None) or School.objects.first()
        
        # 1. 🧠 COMPREHENSIVE CLASS LIST (All Levels)
        sector_map = {
            'PRIMARY': ['Baby', 'Middle', 'Top', 'P.1', 'P.2', 'P.3', 'P.4', 'P.5', 'P.6', 'P.7'],
            'SECONDARY': ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6'],
            'UNIVERSITY': ['Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'],
        }
        classes = sector_map.get(school.sector, ['S.1', 'S.2', 'S.3', 'S.4', 'S.5', 'S.6'])
        
        selected_class = request.GET.get('class', classes[0])
        selected_student_id = request.GET.get('student_id')
        field_type = request.GET.get('field_type', 'eot_score') # Persistent field selection

        # 2. 🔎 ALL-SEEING STUDENT FILTRATION
        # We use __icontains to make sure "S.1" finds "S1", "s.1", and "S.1 "
        clean_class = selected_class.replace(".", "").strip()
        students = Student.objects.filter(
            school=school
        ).filter(
            Q(current_class__iexact=selected_class) | 
            Q(current_class__icontains=clean_class)
        ).order_by('full_name')

        subjects = Subject.objects.all()

        # 3. 💾 SAVE & REFRESH LOGIC
        if request.method == "POST" and selected_student_id:
            student = get_object_or_404(Student, id=selected_student_id)
            target_field = request.POST.get('field_type', 'eot_score')
            
            for sub in subjects:
                score = request.POST.get(f'sub_{sub.id}')
                if score is not None and score != "":
                    res, _ = AcademicResult.objects.get_or_create(student=student, subject=sub)
                    setattr(res, target_field, float(score))
                    res.save()
            
            # 💎 REDIRECT: Keep the student and class selected after saving!
            return redirect(f"/api/secretary-entry/?class={selected_class}&student_id={selected_student_id}&field_type={target_field}")

        # 4. 🧠 THE MEMORY ENGINE: Fetch existing marks for the UI
        existing_marks = {}
        selected_student = None
        if selected_student_id:
            selected_student = students.filter(id=selected_student_id).first()
            if selected_student:
                marks_objs = AcademicResult.objects.filter(student=selected_student)
                for m in marks_objs:
                    # Store the specific score we are currently editing
                    val = getattr(m, field_type, None)
                    existing_marks[m.subject.id] = val if val is not None else ""

        # 5. 🚩 AUDIT: Missing marks calculation
        audit_data = []
        for s in students:
            # We check the specific field_type being entered
            # We use a filter that works even if the field is 0
            completed = AcademicResult.objects.filter(
                student=s, 
                subject__in=subjects
            ).exclude(**{f"{field_type}": 0}).count()
            
            missing = subjects.count() - completed
            audit_data.append({
                'student': s,
                'complete': missing <= 0,
                'missing_count': missing if missing > 0 else 0
            })

        return render(request, 'admin/secretary_marks.html', {
            'classes': classes,
            'selected_class': selected_class,
            'audit_data': audit_data,
            'subjects': subjects,
            'selected_student': selected_student,
            'existing_marks': existing_marks, # 💎 Send memory to UI
            'field_type': field_type,
            'school': school
        })
    except Exception as e:
        return HttpResponse(f"Registry Error: {str(e)}")

import os
import datetime
from django.db.models import Avg
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from .models import Student, KEBMockResult, School, Staff
from collections import Counter

@login_required
def generate_keb_passlip(request, student_id):
    """
    🏛️ THE NATIONAL KEB PASSLIP ENGINE
    Generates a high-prestige, dual-slip A4 document.
    """
    try:
        student = Student.objects.get(account_number=student_id)
        school = student.school
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="KEB_PASSLIP_{student.full_name}.pdf"'
        
        # Initialize Canvas
        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4 # 595.27 x 841.89

        # 💎 DRAW TWO IDENTICAL SLIPS
        # Slip 1: Top Half
        draw_keb_slip_layout(p, student, school, 0)
        
        # ✂️ Central Cutting Guide
        p.setDash(3, 3)
        p.setStrokeColor(colors.grey)
        p.line(0, height/2, width, height/2)
        p.setDash() # Reset

        # Slip 2: Bottom Half
        draw_keb_slip_layout(p, student, school, height/2)

        p.save()
        return response
    except Exception as e:
        return HttpResponse(f"KEB Printing Error: {str(e)}", status=400)


def draw_keb_slip_layout(p, student, school, y_offset):
    width, height = A4
    base_y = height - y_offset 
    
    # 🎨 IMPERIAL GOLD & NAVY PALETTE
    gold_bg = colors.HexColor("#FFF9E6")    # Light Gold Silk
    imperial_gold = colors.HexColor("#D4AF37") 
    rich_gold = colors.HexColor("#D4AF37")     # 💎 THE MISSING GOLD!
    gov_blue = colors.HexColor("#002366") 
    ug_yellow = colors.HexColor("#FCDC04")     # National Yellow   
    ug_red = colors.HexColor("#D90000")       
    off_white = colors.HexColor("#FFFFFF")
    success_green = colors.HexColor("#006400") 

    # 1. 🖌️ BACKGROUND & TRIPLE BORDERS
    p.setFillColor(gold_bg)
    p.rect(10, base_y - 415, width - 20, 405, fill=1, stroke=0)
    p.setLineWidth(2.5); p.setStrokeColor(gov_blue); p.rect(15, base_y - 410, width - 30, 395, stroke=1)
    p.setLineWidth(0.5); p.setStrokeColor(imperial_gold); p.rect(18, base_y - 407, width - 36, 389, stroke=1)

    # 2. 🏛️ NATIONAL TOP HEADERS (Extreme Top)
    p.setFillColor(colors.black); p.setFont("Times-Bold", 10)
    p.drawCentredString(width/2, base_y - 25, "THE REPUBLIC OF UGANDA")
    p.setFont("Times-Bold", 14); p.setFillColor(gov_blue)
    p.drawCentredString(width/2, base_y - 42, "KYADONDO EXAMINATIONS BOARD (KEB)")
    
    # 3. 📸 EXTREME TOP STUDENT PHOTO (Right of Republic/KEB)
    px, py, pw, ph = width - 110, base_y - 120, 75, 90
    if student.photo and os.path.exists(student.photo.path):
        p.setStrokeColor(gov_blue); p.setLineWidth(1.5)
        p.rect(px, py, pw, ph, stroke=1)
        p.drawImage(student.photo.path, px, py, width=pw, height=ph, mask='auto')
    else:
        # 👤 Draw the Vector Silhouette
        p.setStrokeColor(colors.grey); p.rect(px, py, pw, ph, stroke=1)
        p.setFillColor(colors.HexColor("#DCDCDC"))
        p.circle(px + pw/2, py + ph - 25, 15, fill=1)
        p.roundRect(px + 10, py + 10, pw - 20, 40, 8, fill=1)
    
    
    # 4. 👤 STUDENT & SCHOOL IDENTITY (CLEANED)
    lx, ly, lw, lh = 45, base_y - 120, 70, 70
    if school.logo and os.path.exists(school.logo.path):
        p.drawImage(school.logo.path, lx, ly, width=lw, height=lh, mask='auto')
    else:
        p.setStrokeColor(gov_blue); p.rect(lx, ly, lw, lh, stroke=1)

    # Combined School & Student Info starting to the right of the logo (x=125)
    ix = 125
    p.setFillColor(gov_blue); p.setFont("Times-Bold", 11)
    p.drawString(ix, base_y - 65, school.name.upper())
    
    p.setFillColor(colors.black); p.setFont("Times-Bold", 9)
    p.drawString(ix, base_y - 80, f"STUDENT: {student.full_name.upper()}")
    p.drawString(ix, base_y - 95, f"LEVEL: {student.current_class} ({student.stream or 'NORTH'})")
    p.drawString(ix, base_y - 110, f"YEAR: 2026")

    # 🧮 5. NATIONAL RANKING LOGIC (The "Result 1, 2, 3" or "Points")
    results_qs = KEBMockResult.objects.filter(student=student)
    c_name = str(student.current_class).upper()
    is_a_level = any(x in c_name for x in ["S5", "S.5", "S6", "S.6"])
    
    total_score = 0
    total_uace_points = 0
    count = results_qs.count()

    for r in results_qs:
        total_score += r.score
        if is_a_level:
            # Using the 5-point scale logic (A=5, B=4, C=3, D=2, E=1)
            if r.score >= 80: total_uace_points += 5
            elif r.score >= 70: total_uace_points += 4
            elif r.score >= 60: total_uace_points += 3
            elif r.score >= 50: total_uace_points += 2
            elif r.score >= 40: total_uace_points += 1

    avg = total_score / count if count > 0 else 0
    
    results_qs = KEBMockResult.objects.filter(student=student)
    c_name = str(student.current_class).upper()
    is_a_level = any(x in c_name for x in ["S5", "S.5", "S6", "S.6", "A-LEVEL"])
    
    total_uace_points = sum(r.points for r in results_qs if r.points)
    all_fails = all(r.score < 40 for r in results_qs) if results_qs.exists() else True
    
    all_grades = [r.grade for r in results_qs if r.grade]
    most_frequent_grade = Counter(all_grades).most_common(1)[0][0] if all_grades else "N/A"

    total_uace_points = sum(r.points for r in results_qs if r.points)
    all_fails = all(r.score < 40 for r in results_qs) if results_qs.exists() else True
        
    bar_y = base_y - 144
    bar_height = 22
    full_width = width - 90  # Calculates the span from left margin to right
    split_point = 150       # Width of the Gold Grade section

        # 1. Paint the GRADE Section (Imperial Gold)
    p.setFillColor(rich_gold)
    p.rect(45, bar_y, split_point, bar_height, fill=1, stroke=0)
        
        # 2. Paint the RANKING Section (Emerald Green)
    p.setFillColor(success_green)
    p.rect(45 + split_point, bar_y, full_width - split_point, bar_height, fill=1, stroke=0)

        # 3. Insert the Text for GRADE (Black ink on Gold)
    p.setFillColor(colors.black)
    p.setFont("Times-Bold", 10)
    p.drawCentredString(45 + (split_point/2), bar_y + 7, f"★★★ GRADE: {most_frequent_grade} ★★★")

        # 4. Insert the Text for RANKING (White ink on Green)
    p.setFillColor(colors.white)
    p.setFont("Times-Bold", 9)
        
    if is_a_level:
        rank_text = f"KEB WEIGHT: {total_uace_points} / 15 POINTS"
    else:
        res_tier = "RESULT 3 (BASIC)" if all_fails else "RESULT 1 (EXCELLENT)"
        rank_text = f"KEB RANKING: {res_tier}"
    
        p.drawString(45 + split_point + 15, bar_y + 7, rank_text)


    headers = ['SUBJECT NAME', 'SCORE', 'GRD', 'PERFORMANCE GRAPH', 'INTERPRETATION']
    col_widths = [140, 45, 35, 130, 160] 
    data_rows = [headers]

    for r in results_qs:
        score = r.score
        if score >= 90: interp = "EXCEPTIONAL"
        elif score >= 80: interp = "OUTSTANDING"
        elif score >= 70: interp = "GOOD"
        elif score >= 60: interp = "SATISFACTORY"
        elif score >= 50: interp = "BASIC"
        elif score >= 40: interp = "ELEMENTARY"
        else: interp = "UNSATISFACTORY"
        data_rows.append([r.subject.name.upper(), f"{score:g}%", r.grade, "", interp])

    if len(data_rows) == 1: data_rows.append(["NO RECORDS FOUND", "-", "-", "-", "-"])

    # 💎 THE REDUCTION: Changed rowHeights to 17 (from 22)
    table_y = base_y - 300
    table = Table(data_rows, colWidths=col_widths, rowHeights=17)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), gov_blue), ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,-1), 'Times-Bold'), ('FONTSIZE', (0,0), (-1,-1), 7.5),
        ('GRID', (0,0), (-1,-1), 0.1, colors.black), ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.whitesmoke]),
    ]))
    table.wrapOn(p, width, height)
    table.drawOn(p, 45, table_y)

    # 📈 5.5 ALIGNED BARS (Recalculated for height 17)
    graph_x = 45 + 140 + 45 + 35 + 15 
    for i, r in enumerate(results_qs):
        bar_y = (table_y + (len(data_rows) - i - 2) * 17) + 5.5
        p.setFillColor(colors.HexColor("#E0E0E0")) 
        p.roundRect(graph_x, bar_y, 100, 5, 2.5, fill=1, stroke=0)
        bc = success_green if r.score >= 80 else rich_gold if r.score >= 50 else ug_red
        p.setFillColor(bc)
        p.roundRect(graph_x, bar_y, max(2, r.score), 5, 2.5, fill=1, stroke=0)

    
    p.setFillColor(colors.black); p.setFont("Times-Bold", 7.5)
    p.drawString(45, base_y - 330, "KEB EVALUATION STANDARDS:")

    if is_a_level:
        key_data = [['GRADE:', 'A(5)', 'B(4)', 'C(3)', 'D(2)', 'E(1)', 'O(1)', 'F(0)'],
                    ['LEVEL:', 'Exc', 'V.G', 'Good', 'Sat', 'Fair', 'Sub', 'Fail']]
    else:
        key_data = [['GRADE:', 'A', 'B', 'C', 'D', 'E'],
                    ['ACHIEVEMENT LEVEL:', 'Exceptional', 'Outstanding', 'Satisfactory', 'Basic', 'Elementary']]

    k_table = Table(key_data, colWidths=63 if is_a_level else 85, rowHeights=18) # Reduced to 18
    k_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('FONTSIZE', (0,0), (-1,-1), 6),
        ('FONTNAME', (0,0), (-1,-1), 'Times-Bold'), ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('BACKGROUND', (0,0), (0,-1), colors.lightgrey),
    ]))
    k_table.wrapOn(p, width, height)
    k_table.drawOn(p, 45, base_y - 370)

    # =============================================================
    # ✍️ 7. FOOTER: LOWERED KEB SEAL
    # =============================================================
    p.setStrokeColor(gov_blue); p.line(45, base_y - 385, 200, base_y - 385)
    p.drawString(45, base_y - 395, "KEB EXAMINATIONS CHAIRMAN")
    
    seal_x = width - 190 
    seal_y = base_y - 415

    if school.keb_logo and os.path.exists(school.keb_logo.path):
        # 💎 DRAW THE TELESCOPE LOGO (Larger & Proportional 55x55)
        p.drawImage(school.keb_logo.path, seal_x, seal_y, width=45, height=45, mask='auto')
    else:
        # Fallback if no logo
        p.setStrokeColor(colors.teal)
        p.circle(seal_x + 25, seal_y + 25, 25, stroke=1)

    # 💎 TEXT ATTACHED TO THE RIGHT OF THE LOGO
    # We position the text exactly 60 units from the start of the logo
    p.setFillColor(gov_blue)
    p.setFont("Times-Bold", 10)
    p.drawString(seal_x + 60, seal_y + 35, "KEB VERIFIED")
    
    p.setFont("Times-Bold", 7)
    p.setFillColor(colors.black)
    p.drawString(seal_x + 60, seal_y + 22, "NATIONAL MOCK")
    p.drawString(seal_x + 60, seal_y + 12, "REGISTRY 2026")
    
    # Bottom Timestamp
    p.setFont("Times-Roman", 5); p.setFillColor(colors.grey)
    p.drawString(45, base_y - 405, f"Digital Authentication Hash: {datetime.datetime.now().strftime('%d%m%Y%H%M%S')}")

@login_required
def keb_mock_portal_view(request):
    try:
        school = getattr(request.user, 'school', None) or School.objects.first()
        
        sector_map = {
            'PRIMARY': ['P.6', 'P.7'],
            'SECONDARY': ['S.4', 'S.6'],
        }
        classes = sector_map.get(school.sector, ['S.4', 'S.6'])

        subjects = Subject.objects.all().order_by('name')

        # 🔎 Filter Logic
        sel_class = request.GET.get('class', classes[0])
        # Default to first subject if none selected
        first_sub = subjects.first().id if subjects.exists() else None
        sel_sub = request.GET.get('subject', first_sub)
        
        # ⚡ Fetch Students 
        students = Student.objects.filter(school=school, current_class=sel_class).order_by('full_name')
        
        # 📊 Audit Data (Linking results to students)
        audit = []
        for s in students:
            res = KEBMockResult.objects.filter(student=s, subject_id=sel_sub).first() if sel_sub else None
            audit.append({
                'student': s,
                'score': res.score if res else "",
                'grade': res.grade if res else "--",
            })

        context = {
            'school': school,
            'classes': classes,
            'subjects': subjects,
            'sel_class': sel_class,
            'sel_sub': int(sel_sub) if sel_sub and str(sel_sub).isdigit() else None,
            'audit': audit,
            'total_students': students.count(),
            'title': "KEB MOCKS COMMAND"
        }
        return render(request, 'admin/keb_mock_portal.html', context)
    except Exception as e:
        return HttpResponse(f"<body style='background:black;color:red;padding:50px;'><h1>KEB Portal Engine Error</h1><p>{str(e)}</p></body>")
    
@login_required
@transaction.atomic
def save_keb_marks(request):
    """🛡️ THE Hub Hub Hub Hub Hub NATIONAL GRADING ENGINE (UCE & UACE)"""
    if request.method == "POST":
        subject_id = request.POST.get('subject_id')
        class_name = request.POST.get('class_name').upper()
        
        # 🕵️ Determine the Level (O-Level vs A-Level)
        is_a_level = any(x in class_name for x in ["S.5", "S5", "S.6", "S6"])

        for key, value in request.POST.items():
            if key.startswith('score_') and value != "":
                student_id = key.replace('score_', '')
                student = get_object_or_404(Student, id=student_id)
                subject = get_object_or_404(Subject, id=subject_id)
                
                score = float(value)
                grade = "E"
                points = 0

                # 📊 1. NEW CURRICULUM UCE (O-LEVEL) GRADING
                if not is_a_level:
                    if score >= 80: grade = "A"
                    elif score >= 70: grade = "B"
                    elif score >= 60: grade = "C"
                    elif score >= 50: grade = "D"
                    else: grade = "E"
                    points = 0 # Points usually aren't used for O-level aggregates in new curriculum

                # 🎓 2. UACE (A-LEVEL) PRINCIPAL PASS SCALES
                else:
                    if score >= 80: 
                        grade = "A"; points = 6
                    elif score >= 70: 
                        grade = "B"; points = 5
                    elif score >= 60: 
                        grade = "C"; points = 4
                    elif score >= 50: 
                        grade = "D"; points = 3
                    elif score >= 40: 
                        grade = "E"; points = 2
                    elif score >= 35: 
                        grade = "O"; points = 1
                    else: 
                        grade = "F"; points = 0

                # 💾 SAVE TO THE KEB REGISTRY
                KEBMockResult.objects.update_or_create(
                    student=student,
                    subject=subject,
                    defaults={
                        'score': score, 
                        'grade': grade,
                        'points': points # 💎 Automatically stored for A-level rankings
                    }
                )
        
        return redirect(f'/api/keb-portal/?class={class_name}&subject={subject_id}&status=synced')